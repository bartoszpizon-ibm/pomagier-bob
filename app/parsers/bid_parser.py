"""
Parser for a filled-in IBM Special Bid DOCX.
Extracts the key fields needed to generate a Business Justification:
  - client name, model, currency, list price, net price, discount %
  - opportunity context text (Q1 answer)
  - deal background text (Q2 answer)
"""
from __future__ import annotations
import re
from pathlib import Path
from typing import Any

from docx import Document


def parse_bid_docx(source) -> dict[str, Any]:
    """
    Parse a filled Special Bid DOCX file.

    Parameters
    ----------
    source : file-like or Path
        The DOCX file to parse.

    Returns
    -------
    dict with keys:
        client_name, model_name, currency,
        list_price, net_price, discount_pct,
        opportunity_context, deal_background,
        parse_warnings (list[str])
    """
    doc = Document(source)
    result: dict[str, Any] = {
        "client_name": "",
        "model_name": "",
        "currency": "EUR",
        "list_price": 0.0,
        "net_price": 0.0,
        "discount_pct": 0.0,
        "opportunity_context": "",
        "deal_background": "",
        "parse_warnings": [],
    }

    tables = doc.tables
    if len(tables) < 5:
        result["parse_warnings"].append("Unexpected DOCX structure — fewer than 5 tables found.")
        return result

    # ── Table 1 — Channel / Tier ─────────────────────────────────────────
    # row 1 = Tier-1 Distributor, row 2 = Tier-2 Reseller,
    # row 3 = End User (client), row 5 = IBM Sales Rep
    try:
        t1 = tables[1]
        result["client_name"] = _cell_text(t1.rows[3].cells[1])
    except Exception:
        result["parse_warnings"].append("Could not read channel table (table 1).")

    # ── Table 4 — Pricing Summary ────────────────────────────────────────
    # row 1 = End User Buy Price (net), row 2 = Requested Discount %
    try:
        t4 = tables[4]
        net_raw  = _cell_text(t4.rows[1].cells[1])
        disc_raw = _cell_text(t4.rows[2].cells[1])

        # Extract currency and net price — e.g. "1,234,567.00 EUR  (2 × FS9600)"
        curr_m = re.search(r"(EUR|USD|GBP|PLN|CHF)", net_raw, re.IGNORECASE)
        if curr_m:
            result["currency"] = curr_m.group(1).upper()
        net_m = re.search(r"([\d,]+\.?\d*)", net_raw.replace(" ", ""))
        if net_m:
            result["net_price"] = float(net_m.group(1).replace(",", ""))

        # Extract model name from the net price cell parenthetical or full text
        model_m = re.search(r"\(.*?(\bFS\d+|\bESS\d+|\bScale\b[^)]*)\)", net_raw, re.IGNORECASE)
        if model_m:
            result["model_name"] = model_m.group(1).strip()

        # Extract discount %
        disc_m = re.search(r"([\d.]+)\s*%", disc_raw)
        if disc_m:
            result["discount_pct"] = float(disc_m.group(1))
            d = result["discount_pct"] / 100
            if result["net_price"] and d < 1.0:
                result["list_price"] = round(result["net_price"] / (1 - d), 2)
    except Exception as e:
        result["parse_warnings"].append(f"Could not read pricing table (table 4): {e}")

    # ── Table 2 — Q1 Opportunity Context ────────────────────────────────
    try:
        result["opportunity_context"] = _answer_text(tables[2])
    except Exception:
        result["parse_warnings"].append("Could not read opportunity context (table 2).")

    # ── Table 3 — Q2 Deal Background ────────────────────────────────────
    try:
        result["deal_background"] = _answer_text(tables[3])
    except Exception:
        result["parse_warnings"].append("Could not read deal background (table 3).")

    return result


# ── Helpers ──────────────────────────────────────────────────────────────────

def _cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _answer_text(table) -> str:
    """Concatenate all paragraph text from the first cell of a Q/A table,
    skipping the 'Answer:' label itself."""
    lines = []
    for row in table.rows:
        cell = row.cells[0]
        for p in cell.paragraphs:
            t = p.text.strip()
            if t and t.lower().startswith("answer"):
                continue
            if t:
                lines.append(t)
    return "\n".join(lines)
