"""
RFP / RFI Generator for IBM Storage Scale System (3500 / 6000).
Produces a DOCX requirements table tailored to parallel file storage / HPC / AI workloads.
Vendor-neutral wording throughout — no IBM product names in requirement cells.
"""
from __future__ import annotations

import io
import math
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.knowledge.product_db import get_model_info

IBM_BLUE  = RGBColor(0x00, 0x62, 0xFF)
IBM_DARK  = RGBColor(0x16, 0x16, 0x16)
IBM_GRAY  = RGBColor(0x52, 0x52, 0x52)
IBM_LG    = RGBColor(0xF4, 0xF4, 0xF4)
IBM_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ASSETS_DIR = Path(__file__).parent.parent / "assets"
LOGOS_DIR  = ASSETS_DIR / "logos"


# ---------------------------------------------------------------------------
# Number-word helpers
# ---------------------------------------------------------------------------
_PL_WORDS = {1:"jedna",2:"dwie",3:"trzy",4:"cztery",5:"pięć",
             6:"sześć",7:"siedem",8:"osiem",9:"dziewięć",10:"dziesięć"}
_EN_WORDS = {1:"one",2:"two",3:"three",4:"four",5:"five",
             6:"six",7:"seven",8:"eight",9:"nine",10:"ten"}

def _num_word_pl(n): return _PL_WORDS.get(n, str(n))
def _num_word_en(n): return _EN_WORDS.get(n, str(n))
def _floor(v, step=5): return math.floor(v / step) * step


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_scale_rfp(
    project: dict[str, Any],
    client_name: str = "",
    seller_name: str = "",
    iops_override: int | None = None,
    lang: str = "en",
    num_systems: int = 1,
) -> bytes:
    """Generate Storage Scale RFP requirements DOCX and return as bytes."""
    doc = Document()
    _set_page_margins(doc)
    _set_default_font(doc)

    model_code = project.get("model_code", "")
    model_info = get_model_info(model_code)
    iops       = iops_override or project.get("perf_iops_total", 0)

    _add_header_block(doc, project, model_info, client_name, seller_name, lang)
    _add_intro_paragraph(doc, project, model_info, client_name, lang, num_systems=num_systems)
    _add_requirements_table(doc, project, model_info, iops, lang)
    _add_footer_note(doc, lang)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _add_header_block(doc, project, model_info, client_name, seller_name, lang):
    from app.generators.rfp_generator import _add_hrule, _para_space

    if lang == "pl":
        title_txt  = "Specyfikacja Techniczna — Wymagania dla Systemu Pamięci Masowej z Równoległym Systemem Plików"
        client_lbl = "Zamawiający"
        date_lbl   = "Data"
        date_fmt   = "%d.%m.%Y"
    else:
        title_txt  = "Technical Specification — Parallel File Storage System Requirements"
        client_lbl = "Client"
        date_lbl   = "Date"
        date_fmt   = "%B %d, %Y"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_space(title, before=6, after=2)
    run = title.add_run(title_txt)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = IBM_DARK

    sub = doc.add_paragraph()
    _para_space(sub, before=0, after=2)
    run = sub.add_run(
        f"{client_lbl}: {client_name or '—'}  ·  "
        f"{date_lbl}: {date.today().strftime(date_fmt)}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = IBM_GRAY

    _add_hrule(doc)


def _add_intro_paragraph(doc, project, model_info, client_name, lang, num_systems=1):
    from app.generators.rfp_generator import _para_space

    raw_tib    = project.get("raw_tib", 0.0)
    usable_tib = project.get("usable_tib", 0.0)
    num_nodes  = project.get("num_nodes", 1) or 1
    sup        = project.get("support_info") or {}
    sup_years  = sup.get("years", 3)
    sup_cov    = sup.get("coverage", "24×7")
    n          = max(1, int(num_systems))

    raw_tib_f    = _floor(raw_tib)    or round(raw_tib,    1)
    usable_tib_f = _floor(usable_tib) or round(usable_tib, 1)

    if lang == "pl":
        sup_fix = "z fix-time SLA" if sup.get("fix_time") else "bez fix-time SLA"
        _qty = (
            f"Zamówienie obejmuje dostawę {n} (słownie: {_num_word_pl(n)}) systemów "
            f"spełniających poniższe wymagania techniczne. "
            f"Wymagania w tabeli dotyczą parametrów pojedynczego systemu.\n\n"
            if n > 1 else ""
        )
        body = (
            f"{_qty}"
            f"Przedmiotem zamówienia jest dostawa fabrycznie nowego systemu pamięci masowej "
            f"opartego na równoległym systemie plików klasy enterprise, "
            f"wraz z oprogramowaniem zarządzającym i usługami wsparcia technicznego. "
            f"Dostarczone rozwiązanie musi spełniać wszystkie wymagania techniczne określone poniżej.\n\n"
            f"Minimalne wymagane parametry: pojemność fizyczna (raw) {raw_tib_f:.0f} TiB, "
            f"pojemność użyteczna {usable_tib_f:.0f} TiB, "
            f"min. {num_nodes} węzeł(łów) pamięci masowej, "
            f"wsparcie techniczne producenta {sup_cov} przez {sup_years} lat ({sup_fix})."
        )
    else:
        sup_fix = "with fix-time SLA" if sup.get("fix_time") else "without fix-time SLA"
        _qty = (
            f"This procurement covers the delivery of {n} ({_num_word_en(n)}) system(s) "
            f"meeting the technical requirements below. "
            f"All requirements in the table refer to a single system.\n\n"
            if n > 1 else ""
        )
        body = (
            f"{_qty}"
            f"The subject of this procurement is the delivery of a brand-new enterprise-class "
            f"parallel file storage system, together with management software and technical support services. "
            f"The delivered solution must meet all technical requirements specified in the table below.\n\n"
            f"Minimum required parameters: physical (raw) capacity {raw_tib_f:.0f} TiB, "
            f"usable capacity {usable_tib_f:.0f} TiB, "
            f"min. {num_nodes} storage node(s), "
            f"manufacturer support in {sup_cov} mode for {sup_years} years ({sup_fix})."
        )

    p = doc.add_paragraph()
    _para_space(p, before=8, after=6)
    run = p.add_run(body)
    run.font.size = Pt(10)
    run.font.color.rgb = IBM_DARK


def _add_requirements_table(doc, project, model_info, iops, lang):
    from app.generators.rfp_generator import (
        _para_space, _para_space_after_table,
        _set_cell_bg, _set_cell_text,
    )

    rows_data = _build_requirements(project, model_info, iops, lang)

    heading = doc.add_paragraph()
    _para_space(heading, before=8, after=4)
    run = heading.add_run(
        "Tabela Wymagań Technicznych" if lang == "pl" else "Technical Requirements Table"
    )
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = IBM_DARK

    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    col_widths = [Cm(1.2), Cm(4.5), Cm(12.8)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    if lang == "pl":
        h0, h1, h2 = "Lp.", "Nazwa elementu / parametru", "Szczegółowy opis wymagań"
    else:
        h0, h1, h2 = "No.", "Parameter / Element", "Detailed Requirement Description"

    hdr = table.rows[0]
    for cell in hdr.cells:
        _set_cell_bg(cell, IBM_DARK)
    _set_cell_text(hdr.cells[0], h0, bold=True, color=IBM_WHITE, size=9)
    _set_cell_text(hdr.cells[1], h1, bold=True, color=IBM_WHITE, size=9)
    _set_cell_text(hdr.cells[2], h2, bold=True, color=IBM_WHITE, size=9)

    for i, (lp, name, requirement) in enumerate(rows_data):
        row = table.rows[i + 1]
        bg = IBM_LG if i % 2 == 0 else IBM_WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        _set_cell_text(row.cells[0], str(lp),    size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], name,        size=9, bold=True)
        _set_cell_text(row.cells[2], requirement, size=9)

    _para_space_after_table(doc)


def _add_footer_note(doc, lang):
    from app.generators.rfp_generator import _add_hrule, _para_space

    _add_hrule(doc)
    p = doc.add_paragraph()
    _para_space(p, before=4, after=0)
    _date_fmt = "%d.%m.%Y" if lang == "pl" else "%B %d, %Y"
    txt = (
        f"Dokument wygenerowany na podstawie wymagań technicznych · {date.today().strftime(_date_fmt)}"
        if lang == "pl" else
        f"Document generated based on technical requirements · {date.today().strftime(_date_fmt)}"
    )
    run = p.add_run(txt)
    run.font.size = Pt(8)
    run.font.color.rgb = IBM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Requirements builder
# ---------------------------------------------------------------------------

def _build_requirements(project, model_info, iops, lang):
    if lang == "pl":
        return _build_pl(project, model_info, iops)
    return _build_en(project, model_info, iops)


def _build_pl(project, model_info, iops):
    num_nodes     = project.get("num_nodes",     1) or 1
    drives_count  = project.get("drives_count",  0)
    drives_node   = project.get("drives_per_node", drives_count // max(num_nodes, 1))
    raw_tib       = project.get("raw_tib",       0.0)
    raw_tb        = project.get("raw_tb",        0.0)
    usable_tib    = project.get("usable_tib",    0.0)
    usable_tb     = project.get("usable_tb",     0.0)
    raid_type     = project.get("raid_type",     "Erasure Code (8+2p)")
    rebuild_areas = project.get("rebuild_areas", 2)
    network_type  = project.get("network_type",  "InfiniBand NDR 200 Gb/s lub Ethernet 100 GbE")
    network_ports = project.get("network_ports", 2)
    encryption    = project.get("encryption",    False)
    protocols     = project.get("protocol_support", ["NFS", "SMB"])
    proto_str     = ", ".join(protocols) if protocols else "NFS, SMB"
    fs_type       = project.get("filesystem_type", "równoległy system plików klasy enterprise (POSIX)")

    raw_tib_f    = _floor(raw_tib)    or round(raw_tib,    1)
    raw_tb_f     = _floor(raw_tb)     or round(raw_tb,     1)
    usable_tib_f = _floor(usable_tib) or round(usable_tib, 1)
    usable_tb_f  = _floor(usable_tb)  or round(usable_tb,  1)

    enc_txt = (
        "Wymagane — szyfrowanie danych at-rest bez wpływu na wydajność systemu; obsługa KMIP"
        if encryption else
        "Wymagane — należy dostarczyć możliwość aktywacji szyfrowania; obsługa serwerów KMIP"
    )

    sup       = project.get("support_info") or {}
    sup_years = sup.get("years",    3)
    sup_cov   = sup.get("coverage", "24×7")
    sup_fix   = "z fix-time SLA" if sup.get("fix_time") else "bez fix-time SLA"

    tp_mib  = project.get("perf_throughput_mib", 0.0)
    tp_floor = (int(tp_mib / 100) * 100) if tp_mib else 0
    iops_f   = (int(iops / 25_000) * 25_000) if iops else 0

    if tp_floor:
        perf_txt = (
            f"System musi zapewniać przepustowość sekwencyjną: min. {tp_floor:,} MiB/s "
            f"(odczyt sekwencyjny 1 MiB, losowy dostęp wieloklientowy)."
        )
        if iops_f:
            perf_txt += f" Losowa wydajność: min. {iops_f:,} IOPS (4 KiB, mieszany ruch)."
    elif iops_f:
        perf_txt = f"System musi zapewniać wydajność min. {iops_f:,} IOPS dla skonfigurowanego profilu obciążenia."
    else:
        perf_txt = (
            "System musi zapewniać wydajność odpowiednią do obciążeń AI/HPC. "
            "Szczegółowe wymagania przepustowości należy uzupełnić na podstawie analizy workload."
        )

    return [
        (1,  "Architektura systemu",
         f"System musi być oparty na architekturze węzłowej (scale-out), składającej się z minimum {num_nodes} węzła(łów) "
         f"pamięci masowej umożliwiających dodawanie kolejnych węzłów bez przerwy w pracy (non-disruptive). "
         f"Wszystkie węzły muszą być montowane w standardowej szafie 19\"."),

        (2,  "Równoległy system plików",
         f"System musi wykorzystywać {fs_type} zapewniający jednolitą globalną przestrzeń nazw (global namespace) "
         f"dostępną jednocześnie przez wszystkie węzły obliczeniowe. "
         f"System plików musi być zgodny ze standardem POSIX i obsługiwać min. {proto_str}. "
         f"Wymagana natywna obsługa protokołów dostępu bez dodatkowych bram (gateway) dla NFS v3/v4."),

        (3,  "Pojemność",
         f"Całkowita pojemność fizyczna (raw): min. {raw_tib_f:.0f} TiB ({raw_tb_f:.0f} TB). "
         f"Pojemność użyteczna: min. {usable_tib_f:.0f} TiB ({usable_tb_f:.0f} TB) "
         f"przy zastosowaniu mechanizmu ochrony {raid_type}. "
         f"Minimalna liczba napędów NVMe: {drives_count} szt. ({drives_node} na węzeł). "
         f"Zamawiający nie dopuszcza dysków rotacyjnych jako warstwy podstawowej."),

        (4,  "Ochrona danych",
         f"System musi stosować mechanizm kodowania korekcyjnego (erasure coding) — "
         f"ochrona {raid_type} z odpornością na jednoczesną awarię minimum 2 węzłów lub napędów. "
         f"Przestrzeń zapasowa musi być rozproszona równomiernie we wszystkich węzłach klastra "
         f"(bez dedykowanych fizycznych węzłów spare). "
         f"Minimalna liczba domen awarii (rebuild areas/fault domains): {rebuild_areas}."),

        (5,  "Wydajność",
         perf_txt),

        (6,  "Interfejsy sieciowe",
         f"Każdy węzeł musi być wyposażony w minimum {network_ports} porty sieci danych "
         f"({network_type}) do komunikacji klient-serwer (data path). "
         f"Wymagana redundancja połączeń sieciowych — brak pojedynczego punktu awarii (no SPOF). "
         f"System musi obsługiwać protokoły RDMA (Remote Direct Memory Access) w celu minimalizacji latencji."),

        (7,  "Szyfrowanie i bezpieczeństwo",
         f"Szyfrowanie danych: {enc_txt}. "
         f"Wymagana kontrola dostępu oparta na rolach (RBAC) dla interfejsu zarządzania. "
         f"Obsługa uwierzytelniania przez LDAP/Active Directory. "
         f"Immutable snapshots — kopie niemutowalne chroniące przed ransomware i przypadkowym usunięciem."),

        (8,  "Skalowalność",
         f"System musi umożliwiać rozbudowę pojemności i wydajności przez dodawanie węzłów "
         f"bez przerwy w pracy (non-disruptive) i bez rekonfiguracji istniejących klientów. "
         f"Skalowanie musi być liniowe: dodanie węzła proporcjonalnie zwiększa zarówno pojemność, jak i przepustowość. "
         f"Docelowa pojemność klastra: min. kilka PiB (bez przebudowy architektury)."),

        (9,  "Migawki i replikacja",
         f"System musi obsługiwać migawki spójne na poziomie systemu plików (filesystem-level snapshots) "
         f"z możliwością tworzenia migawek niezmiennych (immutable) z wymuszanym czasem retencji. "
         f"Wymagana replikacja asynchroniczna między lokalizacjami (site-to-site) "
         f"z granularnością na poziomie katalogu lub systemu plików."),

        (10, "Zarządzanie",
         f"System musi być zarządzany za pomocą interfejsu graficznego (GUI) dostępnego przez przeglądarkę webową "
         f"oraz interfejsu tekstowego (CLI) i REST API. "
         f"Wymagana integracja z systemami monitorowania (SNMP, Syslog, Prometheus/Grafana). "
         f"Zarządzanie policyką danych: automatyczny tiering na podstawie częstotliwości dostępu "
         f"(hot/warm/cold data placement). "
         f"Obsługa wielodostępu (multi-tenancy) z izolacją przestrzeni nazw (fileset-level quotas)."),

        (11, "Integracja z obciążeniami AI/HPC",
         f"System musi posiadać natywne integracje z popularnymi frameworkami AI/ML "
         f"(np. scheduler MPI, obsługa metadanych dla treningu modeli). "
         f"Wymagana obsługa protokołu S3-compatible object storage jako warstwy tiering. "
         f"Kompatybilność z konteneryzacją (CSI driver dla Kubernetes/OpenShift) musi być udokumentowana."),

        (12, "Niezawodność i dostępność",
         f"Wszystkie krytyczne komponenty (kontrolery węzłów, zasilacze, wentylatory, interfejsy sieciowe) "
         f"muszą być nadmiarowe i wymienialne w trakcie pracy systemu (hot-swap). "
         f"System musi tolerować jednoczesną awarię min. 2 napędów bez utraty danych i bez przerwy w dostępie. "
         f"Gwarantowany czas MTTRu po awarii napędu (rebuild): dane muszą być odtworzone w czasie < 2h na TB."),

        (13, "Środowisko fizyczne",
         f"Montaż w standardowej szafie rack 19\". "
         f"Zasilanie: dwa niezależne zasilacze hot-swap (200–240V, 50/60 Hz). "
         f"Chłodzenie: przez przedni wlot / tylny wylot powietrza (standardowy przepływ szafy serwerowej). "
         f"Temperatura pracy: 10–35°C (ASHRAE A2)."),

        (14, "Wsparcie techniczne",
         f"Producent musi zapewniać wsparcie techniczne w trybie {sup_cov} przez minimum {sup_years} lat "
         f"od daty odbioru ({sup_fix}). "
         f"Wymagane: centralne narzędzie monitorowania proaktywnego z alertowaniem online. "
         f"Możliwość zdalnej diagnostyki i rozwiązywania problemów bez konieczności wizyty serwisowej."),
    ]


def _build_en(project, model_info, iops):
    num_nodes     = project.get("num_nodes",     1) or 1
    drives_count  = project.get("drives_count",  0)
    drives_node   = project.get("drives_per_node", drives_count // max(num_nodes, 1))
    raw_tib       = project.get("raw_tib",       0.0)
    raw_tb        = project.get("raw_tb",        0.0)
    usable_tib    = project.get("usable_tib",    0.0)
    usable_tb     = project.get("usable_tb",     0.0)
    raid_type     = project.get("raid_type",     "Erasure Code (8+2p)")
    rebuild_areas = project.get("rebuild_areas", 2)
    network_type  = project.get("network_type",  "InfiniBand NDR 200 Gb/s or Ethernet 100 GbE")
    network_ports = project.get("network_ports", 2)
    encryption    = project.get("encryption",    False)
    protocols     = project.get("protocol_support", ["NFS", "SMB"])
    proto_str     = ", ".join(protocols) if protocols else "NFS, SMB"
    fs_type       = project.get("filesystem_type", "enterprise-class parallel file system (POSIX)")

    raw_tib_f    = _floor(raw_tib)    or round(raw_tib,    1)
    raw_tb_f     = _floor(raw_tb)     or round(raw_tb,     1)
    usable_tib_f = _floor(usable_tib) or round(usable_tib, 1)
    usable_tb_f  = _floor(usable_tb)  or round(usable_tb,  1)

    enc_txt = (
        "Required — data-at-rest encryption with no performance impact; KMIP key server support"
        if encryption else
        "Required — encryption activation capability must be included; KMIP key server support"
    )

    sup       = project.get("support_info") or {}
    sup_years = sup.get("years",    3)
    sup_cov   = sup.get("coverage", "24×7")
    sup_fix   = "with fix-time SLA" if sup.get("fix_time") else "without fix-time SLA"

    tp_mib   = project.get("perf_throughput_mib", 0.0)
    tp_floor = (int(tp_mib / 100) * 100) if tp_mib else 0
    iops_f   = (int(iops / 25_000) * 25_000) if iops else 0

    if tp_floor:
        perf_txt = (
            f"The system must deliver sequential throughput of at least {tp_floor:,} MiB/s "
            f"(1 MiB sequential read, multi-client concurrent access)."
        )
        if iops_f:
            perf_txt += f" Random I/O: min. {iops_f:,} IOPS (4 KiB mixed workload)."
    elif iops_f:
        perf_txt = f"The system must deliver at least {iops_f:,} IOPS for the configured workload profile."
    else:
        perf_txt = (
            "The system must deliver throughput appropriate for AI/HPC workloads. "
            "Detailed throughput requirements shall be defined based on workload analysis."
        )

    return [
        (1,  "System Architecture",
         f"The system must be based on a scale-out node architecture consisting of a minimum of {num_nodes} storage node(s) "
         f"that allow non-disruptive addition of further nodes. "
         f"All nodes must be rack-mountable in a standard 19″ cabinet."),

        (2,  "Parallel File System",
         f"The system must use {fs_type} providing a unified global namespace "
         f"accessible simultaneously by all compute clients. "
         f"The file system must be POSIX-compliant and natively support {proto_str}. "
         f"Native protocol support (NFS v3/v4) without requiring additional gateways is mandatory."),

        (3,  "Capacity",
         f"Total physical (raw) capacity: min. {raw_tib_f:.0f} TiB ({raw_tb_f:.0f} TB). "
         f"Usable capacity: min. {usable_tib_f:.0f} TiB ({usable_tb_f:.0f} TB) "
         f"with {raid_type} data protection. "
         f"Minimum NVMe drive count: {drives_count} ({drives_node} per node). "
         f"Rotational (HDD) drives are not permitted as the primary storage tier."),

        (4,  "Data Protection",
         f"The system must use erasure coding — {raid_type} protection "
         f"tolerating simultaneous failure of at least 2 nodes or drives. "
         f"Spare capacity must be distributed evenly across all nodes in the cluster "
         f"(no dedicated physical spare nodes or drives). "
         f"Minimum number of fault domains (rebuild areas): {rebuild_areas}."),

        (5,  "Performance",
         perf_txt),

        (6,  "Network Interfaces",
         f"Each storage node must be equipped with a minimum of {network_ports} data network ports "
         f"({network_type}) for client-to-storage data path. "
         f"Network redundancy is mandatory — no single point of failure. "
         f"The system must support RDMA (Remote Direct Memory Access) to minimise access latency."),

        (7,  "Encryption & Security",
         f"Data encryption: {enc_txt}. "
         f"Role-based access control (RBAC) for the management interface is required. "
         f"LDAP/Active Directory authentication integration is required. "
         f"Immutable snapshots for ransomware and accidental-deletion protection must be supported."),

        (8,  "Scalability",
         f"The system must allow non-disruptive capacity and performance expansion "
         f"by adding nodes without reconfiguring existing clients. "
         f"Scaling must be linear — adding a node must proportionally increase both capacity and throughput. "
         f"Target cluster capacity must reach multi-PiB scale without architectural redesign."),

        (9,  "Snapshots & Replication",
         f"The system must support file-system-level consistent snapshots "
         f"with the ability to enforce immutable (WORM) retention policies. "
         f"Asynchronous site-to-site replication at directory or file system granularity is required."),

        (10, "Management",
         f"The system must be managed via a browser-based GUI, CLI, and REST API. "
         f"Integration with monitoring systems (SNMP, Syslog, Prometheus/Grafana) is required. "
         f"Data lifecycle policy management with automated hot/warm/cold tiering must be supported. "
         f"Multi-tenancy with fileset-level namespace isolation and quota enforcement is required."),

        (11, "AI/HPC Workload Integration",
         f"The system must provide native integrations with AI/ML training frameworks and HPC schedulers. "
         f"S3-compatible object storage protocol for data tiering must be supported. "
         f"A documented Kubernetes/OpenShift CSI driver must be available."),

        (12, "Reliability & Availability",
         f"All critical components (node controllers, PSUs, fans, network interfaces) "
         f"must be redundant and hot-swappable. "
         f"The system must tolerate simultaneous failure of at least 2 drives without data loss or access interruption. "
         f"Drive rebuild must complete within 2 hours per TB."),

        (13, "Physical Environment",
         f"Standard 19″ rack-mountable. "
         f"Power: dual redundant hot-swap PSUs (200–240V, 50/60 Hz). "
         f"Cooling: front-to-rear airflow (standard server rack orientation). "
         f"Operating temperature: 10–35°C (ASHRAE A2)."),

        (14, "Technical Support",
         f"The manufacturer must provide {sup_cov} technical support for a minimum of {sup_years} years "
         f"from the acceptance date ({sup_fix}). "
         f"Proactive monitoring with online alerting is required. "
         f"Remote diagnostics and issue resolution without on-site visits must be available."),
    ]


# ---------------------------------------------------------------------------
# DOCX helpers (reuse from FlashSystem rfp_generator)
# ---------------------------------------------------------------------------

def _set_page_margins(doc):
    from app.generators.rfp_generator import _set_page_margins as _fs
    _fs(doc)

def _set_default_font(doc):
    from app.generators.rfp_generator import _set_default_font as _fs
    _fs(doc)
