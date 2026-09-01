"""
RFP / RFI Generator — Polish Technical Requirements Table.
Produces a DOCX with a 17-row requirements table auto-filled from parsed project data.
"""

from __future__ import annotations

import io
import math
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

from ..knowledge.product_db import get_model_info

# IBM design tokens
IBM_BLUE   = RGBColor(0x00, 0x62, 0xFF)
IBM_DARK   = RGBColor(0x16, 0x16, 0x16)
IBM_GRAY   = RGBColor(0x52, 0x52, 0x52)
IBM_LG     = RGBColor(0xF4, 0xF4, 0xF4)
IBM_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
IBM_GREEN  = RGBColor(0x19, 0x8A, 0x3E)

ASSETS_DIR = Path(__file__).parent.parent / "assets"
LOGOS_DIR  = ASSETS_DIR / "logos"
IMAGES_DIR = ASSETS_DIR / "images"


# ---------------------------------------------------------------------------
# Number-word helpers (used in intro paragraph for quantity of systems)
# ---------------------------------------------------------------------------
_PL_WORDS = {
    1: "jedna", 2: "dwie", 3: "trzy", 4: "cztery", 5: "pięć",
    6: "sześć", 7: "siedem", 8: "osiem", 9: "dziewięć", 10: "dziesięć",
}
_EN_WORDS = {
    1: "one", 2: "two", 3: "three", 4: "four", 5: "five",
    6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten",
}

def _num_word_pl(n: int) -> str:
    return _PL_WORDS.get(n, str(n))

def _num_word_en(n: int) -> str:
    return _EN_WORDS.get(n, str(n))


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_rfp(
    project: dict[str, Any],
    client_name: str = "",
    seller_name: str = "",
    iops_override: int | None = None,
    lang: str = "en",
    num_systems: int = 1,
) -> bytes:
    """Generate RFP requirements DOCX and return as bytes."""

    doc = Document()
    _set_page_margins(doc)
    _set_default_font(doc)

    model_code = project.get("model_code", "")
    model_info = get_model_info(model_code)
    iops = iops_override or project.get("perf_iops_total", 0)

    _add_header_block(doc, project, model_info, client_name, seller_name, lang)
    _add_intro_paragraph(doc, project, model_info, client_name, lang, num_systems=num_systems)
    _add_requirements_table(doc, project, model_info, iops, lang)
    _add_footer_note(doc, lang)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Capacity helpers & drive-slot table (used in both intro and requirements)
# ---------------------------------------------------------------------------

# Drive-slot maximums per model family (short name → max physical slots)
# Source: IBM FlashSystem Sales Manual
_MAX_DRIVE_SLOTS: dict[str, int] = {
    "FS5600": 12, "FS5200": 12,          # 1U, 12 drive slots
    "FS5045": 24, "FS5015": 24,          # 2U, 24 drive slots
    "FS7600": 32, "FS7300": 24,          # FS7600=32, FS7300=24
    "FS9600": 32, "FS9500": 48,          # FS9600=32, FS9500=48
    "FSC200": 24,
}

# Controller CPU core counts per model family (short name → total cores across both nodes/controllers)
# Source: IBM FlashSystem Sales Manual specifications
_CONTROLLER_CORES: dict[str, tuple[int, str]] = {
    # (total_cores, description)
    "FS9600": (96, "2 nodes × 1 CPU per node × 48 cores = 96 cores total"),
    "FS9500": (96, "2 nodes × 1 CPU per node × 48 cores = 96 cores total"),
    "FS7600": (32, "2 nodes × 1 CPU per node × 16 cores = 32 cores total"),
    "FS7300": (32, "2 nodes × 1 CPU per node × 16 cores = 32 cores total"),
    "FS5600": (24, "2 nodes × 1 CPU per node × 12 cores = 24 cores total"),
    "FS5200": (24, "2 nodes × 1 CPU per node × 12 cores = 24 cores total"),
    "FS5045": (24, "2 nodes × 1 CPU per node × 12 cores = 24 cores total"),
    "FS5015": (16, "2 nodes × 1 CPU per node × 8 cores = 16 cores total"),
    "FSC200": (40, "2 nodes × 2 CPUs per node × 10 cores = 40 cores total"),
}


def _floor_tib(value: float, step: int = 5) -> float:
    """Round down to nearest multiple of step (default 5 TiB / 5 TB)."""
    return math.floor(value / step) * step


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def _add_header_block(doc, project, model_info, client_name, seller_name, lang="en"):
    """Cover block: logo + title."""
    if lang == "pl":
        title_txt = "Specyfikacja Techniczna — Wymagania dla Macierzy Dyskowej"
        client_lbl = "Zamawiający"
        date_lbl   = "Data"
        date_fmt   = "%d.%m.%Y"
    else:
        title_txt = "Technical Specification — Enterprise Storage Requirements"
        client_lbl = "Client"
        date_lbl   = "Date"
        date_fmt   = "%B %d, %Y"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_space(title, before=6, after=2)
    run = title.add_run(title_txt)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = IBM_DARK

    sub = doc.add_paragraph()
    _para_space(sub, before=0, after=2)
    run = sub.add_run(
        f"{client_lbl}: {client_name or '—'}  ·  "
        f"{date_lbl}: {date.today().strftime(date_fmt)}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = IBM_GRAY

    _add_hrule(doc)


def _add_intro_paragraph(doc, project, model_info, client_name, lang="en", num_systems: int = 1):
    """Neutral intro paragraph — no product name or manufacturer."""
    raw_tib    = project.get("raw_tib", 0.0)
    usable_tib = project.get("usable_tib", 0.0)
    cache_gb   = project.get("cache_gb", 256)
    raid_type  = project.get("raid_type", "RAID 6")
    sup        = project.get("support_info") or {}
    sup_years  = sup.get("years", 3)
    sup_coverage = sup.get("coverage", "24×7")

    raw_tib_floor    = _floor_tib(raw_tib)    or round(raw_tib,    1)
    usable_tib_floor = _floor_tib(usable_tib) or round(usable_tib, 1)
    n = max(1, int(num_systems))

    if lang == "pl":
        sup_fix = "z fix-time SLA" if sup.get("fix_time") else "bez fix-time SLA"
        _qty_note = (
            f"Zamówienie obejmuje dostawę {n} (słownie: {_num_word_pl(n)}) macierzy dyskowych "
            f"spełniających poniższe wymagania techniczne. "
            f"Wymagania techniczne w tabeli dotyczą parametrów pojedynczej macierzy.\n\n"
            if n > 1 else ""
        )
        body_txt = (
            f"{_qty_note}"
            f"Przedmiotem zamówienia jest dostawa fabrycznie nowej macierzy dyskowej klasy enterprise "
            f"wraz z oprogramowaniem zarządzającym i usługami wsparcia technicznego. "
            f"Dostarczone rozwiązanie musi spełniać wszystkie wymagania techniczne określone w poniższej tabeli.\n\n"
            f"Minimalne wymagane parametry techniczne: pojemność fizyczna (raw) {raw_tib_floor:.0f} TiB, "
            f"pojemność użyteczna {usable_tib_floor:.0f} TiB ({raid_type}), "
            f"pamięć cache min. {cache_gb} GB, "
            f"wsparcie techniczne producenta w trybie {sup_coverage} przez {sup_years} lat ({sup_fix})."
        )
    else:
        sup_fix = "with fix-time SLA" if sup.get("fix_time") else "without fix-time SLA"
        _qty_note = (
            f"This procurement covers the delivery of {n} ({_num_word_en(n)}) storage arrays "
            f"meeting the technical requirements below. "
            f"All technical requirements in the table refer to a single array unit.\n\n"
            if n > 1 else ""
        )
        body_txt = (
            f"{_qty_note}"
            f"The subject of this procurement is the delivery of a brand-new enterprise-class disk array "
            f"together with management software and technical support services. "
            f"The delivered solution must meet all technical requirements specified in the table below.\n\n"
            f"Minimum required technical parameters: physical (raw) capacity {raw_tib_floor:.0f} TiB, "
            f"usable capacity {usable_tib_floor:.0f} TiB ({raid_type}), "
            f"cache memory min. {cache_gb} GB, "
            f"manufacturer support in {sup_coverage} mode for {sup_years} years ({sup_fix})."
        )

    p = doc.add_paragraph()
    _para_space(p, before=8, after=6)
    run = p.add_run(body_txt)
    run.font.size = Pt(10)
    run.font.color.rgb = IBM_DARK


def _add_requirements_table(doc, project, model_info, iops: int, lang="en"):
    """Main requirements table (3 columns)."""
    rows_data = _build_requirements(project, model_info, iops, lang)

    # Table heading
    heading = doc.add_paragraph()
    _para_space(heading, before=8, after=4)
    run = heading.add_run(
        "Tabela Wymagań Technicznych" if lang == "pl" else "Technical Requirements Table"
    )
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = IBM_DARK

    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Column widths
    col_widths = [Cm(1.2), Cm(4.5), Cm(12.8)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]

    # Header row
    if lang == "pl":
        h0, h1, h2 = "Lp.", "Nazwa elementu / parametru", "Szczegółowy opis wymagań"
    else:
        h0, h1, h2 = "No.", "Parameter / Element", "Detailed Requirement Description"
    hdr = table.rows[0]
    _set_cell_bg(hdr.cells[0], IBM_DARK)
    _set_cell_bg(hdr.cells[1], IBM_DARK)
    _set_cell_bg(hdr.cells[2], IBM_DARK)
    _set_cell_text(hdr.cells[0], h0, bold=True, color=IBM_WHITE, size=9)
    _set_cell_text(hdr.cells[1], h1, bold=True, color=IBM_WHITE, size=9)
    _set_cell_text(hdr.cells[2], h2, bold=True, color=IBM_WHITE, size=9)

    # Data rows
    for i, (lp, name, requirement) in enumerate(rows_data):
        row = table.rows[i + 1]
        bg = IBM_LG if i % 2 == 0 else IBM_WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)
        _set_cell_text(row.cells[0], str(lp),  size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], name,      size=9, bold=True)
        _set_cell_text(row.cells[2], requirement, size=9)

    _para_space_after_table(doc)


def _add_footer_note(doc, lang="en"):
    _add_hrule(doc)
    p = doc.add_paragraph()
    _para_space(p, before=4, after=0)
    _date_fmt = "%d.%m.%Y" if lang == "pl" else "%B %d, %Y"
    if lang == "pl":
        footer_txt = f"Dokument wygenerowany na podstawie wymagań technicznych · {date.today().strftime(_date_fmt)}"
    else:
        footer_txt = f"Document generated based on technical requirements · {date.today().strftime(_date_fmt)}"
    run = p.add_run(footer_txt)
    run.font.size = Pt(8)
    run.font.color.rgb = IBM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Requirements content builder
# ---------------------------------------------------------------------------

def _build_requirements(project: dict, model_info: dict, iops: int, lang: str = "en") -> list[tuple]:
    """Return list of (lp, name, requirement_text) tuples for 18 rows.
    Vendor-neutral — no IBM/FlashSystem references in the output text.
    """
    if lang == "pl":
        return _build_requirements_pl(project, model_info, iops)
    return _build_requirements_en(project, model_info, iops)


def _build_requirements_pl(project: dict, model_info: dict, iops: int) -> list[tuple]:
    """Polish version of requirements."""

    is_hybrid     = project.get("is_hybrid", False)
    drives_count  = project.get("drives_count", 0)
    raw_tib       = project.get("raw_tib",   0.0)
    raw_tb        = project.get("raw_tb",    0.0)
    usable_tib    = project.get("usable_tib",  0.0)
    usable_tb     = project.get("usable_tb",  0.0)
    cache_gb      = project.get("cache_gb",   256)
    fc_ports      = project.get("fc_ports",   8)
    raid_type     = project.get("raid_type",  "RAID 6")
    rebuild_areas = project.get("rebuild_areas", 2)
    form_factor   = model_info.get("form_factor", "1U")
    rack_units    = project.get("rack_units", None)
    encryption    = project.get("encryption", False)
    io_groups     = project.get("io_groups", 2)
    cable_qty     = project.get("cable_qty", 0)
    cable_label   = f"{cable_qty} szt. kabli FC" if cable_qty else "kable FC zgodne z liczbą portów"

    # Hybrid tier fields
    nvme_drives_count  = project.get("nvme_drives_count",  drives_count)
    nvme_raw_tib       = project.get("nvme_raw_tib",       0.0)
    nvme_raw_tb        = project.get("nvme_raw_tb",        0.0)
    nvme_usable_tib    = project.get("nvme_usable_tib",    0.0)
    nvme_usable_tb     = project.get("nvme_usable_tb",     0.0)
    drive_type         = project.get("drive_type",         "")
    hdd_drives_count   = project.get("hdd_drives_count",   0)
    hdd_raw_tib        = project.get("hdd_raw_tib",        0.0)
    hdd_raw_tb         = project.get("hdd_raw_tb",         0.0)
    hdd_usable_tib     = project.get("hdd_usable_tib",     0.0)
    hdd_usable_tb      = project.get("hdd_usable_tb",      0.0)
    hdd_drive_type     = project.get("hdd_drive_type",     "")
    hdd_enclosure      = project.get("hdd_enclosure",      "")
    hdd_enclosure_qty  = project.get("hdd_enclosure_qty",  1)

    # Max drive slots — use model-specific value or fall back to drives_count
    model_short   = model_info.get("short", "")
    max_slots     = _MAX_DRIVE_SLOTS.get(model_short, max(drives_count, 12))
    _cores_total, _cores_desc = _CONTROLLER_CORES.get(model_short, (24, "min. 24 rdzeni na parę kontrolerów"))

    # Rounded-down capacities (step 5 TiB / 5 TB)
    raw_tib_floor    = _floor_tib(raw_tib)    or round(raw_tib,    1)
    raw_tb_floor     = _floor_tib(raw_tb)     or round(raw_tb,     1)
    usable_tib_floor = _floor_tib(usable_tib) or round(usable_tib, 1)
    usable_tb_floor  = _floor_tib(usable_tb)  or round(usable_tb,  1)

    # Encryption — vendor-neutral wording
    enc_txt = (
        "Wymagane — szyfrowanie danych at-rest bez wpływu na wydajność systemu; "
        "obsługa zewnętrznych serwerów kluczy (KMIP) oraz lokalnych kluczy USB"
        if encryption else
        "Wymagane — należy dostarczyć licencję szyfrowania; obsługa serwerów KMIP"
    )

    sup          = project.get("support_info") or {}
    sup_years    = sup.get("years",    3)
    sup_coverage = sup.get("coverage", "24×7")
    sup_fix      = "z fix-time SLA" if sup.get("fix_time") else "bez fix-time SLA"

    perf_iops_sub1 = project.get("perf_iops_max_sub1ms", 0)
    throughput_mib = project.get("perf_throughput_mib", 0.0)

    # Round down to nearest 25 000 IOPS and 100 MiB/s to avoid pinpointing exact vendor values
    iops_sub1_floor = (int(perf_iops_sub1 / 25_000) * 25_000) if perf_iops_sub1 else 0
    tp_floor        = (int(throughput_mib / 100) * 100)        if throughput_mib  else 0

    if iops_sub1_floor:
        perf_txt = (
            f"System musi zapewniać wydajność co najmniej {iops_sub1_floor:,} IOPS "
            f"przy latencji poniżej 1 ms (pomiar dla skonfigurowanego profilu obciążenia)."
        )
        if tp_floor:
            perf_txt += f" Wymagana przepustowość sekwencyjna: min. {tp_floor:,} MiB/s."
    elif iops:
        perf_txt = (
            f"System musi zapewniać wydajność min. {iops:,} IOPS dla skonfigurowanego profilu "
            f"obciążenia (blok 4 KiB, odczyt/zapis mieszany)."
        )
        if tp_floor:
            perf_txt += f" Przepustowość sekwencyjna: min. {tp_floor:,} MiB/s."
    else:
        perf_txt = (
            "System musi zapewniać wydajność odpowiednią do wymagań Zamawiającego. "
            "Szczegółowe wymagania IOPS należy uzupełnić na podstawie analizy obciążenia."
        )

    # --- Enclosure text (PL) ---
    if is_hybrid and hdd_enclosure:
        _enc_txt = (
            f"Macierz musi umożliwiać instalację w standardowej szafie serwerowej RACK 19\". "
            f"Rozwiązanie składa się z obudowy kontrolerowej ({form_factor}) "
            f"oraz {hdd_enclosure_qty} obudowy rozszerzającej {hdd_enclosure} "
            f"z dyskami HDD NL-SAS. "
            f"Łączna zajętość w szafie RACK: {rack_units} U. "
            f"Obudowa kontrolerowa musi zawierać co najmniej dwa kontrolery "
            f"pracujące w trybie active-active."
        )
    else:
        _enc_txt = (
            f"Macierz musi umożliwiać instalację w standardowej szafie serwerowej RACK 19\". "
            f"Obudowa podstawowa o wysokości co najwyżej {form_factor} musi zawierać "
            f"co najmniej dwa kontrolery pracujące w trybie active-active "
            f"oraz {max_slots} slotów na nośniki danych."
        )

    # --- Capacity text (PL) ---
    if is_hybrid and hdd_drives_count > 0:
        _nvme_raw_tib_fl  = _floor_tib(nvme_raw_tib)    or round(nvme_raw_tib,  1)
        _nvme_raw_tb_fl   = _floor_tib(nvme_raw_tb)     or round(nvme_raw_tb,   1)
        _nvme_us_tib_fl   = _floor_tib(nvme_usable_tib) or round(nvme_usable_tib, 1)
        _nvme_us_tb_fl    = _floor_tib(nvme_usable_tb)  or round(nvme_usable_tb,  1)
        _hdd_raw_tib_fl   = _floor_tib(hdd_raw_tib)     or round(hdd_raw_tib,   1)
        _hdd_raw_tb_fl    = _floor_tib(hdd_raw_tb)      or round(hdd_raw_tb,    1)
        _hdd_us_tib_fl    = _floor_tib(hdd_usable_tib)  or round(hdd_usable_tib, 1)
        _hdd_us_tb_fl     = _floor_tib(hdd_usable_tb)   or round(hdd_usable_tb,  1)
        _enc_str          = f" w obudowie {hdd_enclosure}" if hdd_enclosure else ""
        _cap_txt = (
            f"System musi posiadać dwuwarstwową architekturę pojemnościową (Tier NVMe + Tier HDD). "
            f"Tier NVMe: {nvme_drives_count} nośników NVMe, "
            f"pojemność fizyczna min. {_nvme_raw_tib_fl:.0f} TiB ({_nvme_raw_tb_fl:.0f} TB), "
            f"pojemność użyteczna min. {_nvme_us_tib_fl:.0f} TiB ({_nvme_us_tb_fl:.0f} TB). "
            f"Tier HDD: {hdd_drives_count} dysków NL-SAS{_enc_str}, "
            f"pojemność fizyczna min. {_hdd_raw_tib_fl:.0f} TiB ({_hdd_raw_tb_fl:.0f} TB), "
            f"pojemność użyteczna min. {_hdd_us_tib_fl:.0f} TiB ({_hdd_us_tb_fl:.0f} TB). "
            f"Łączna pojemność fizyczna (raw): min. {raw_tib_floor:.0f} TiB ({raw_tb_floor:.0f} TB). "
            f"Łączna pojemność użyteczna ({raid_type}): min. {usable_tib_floor:.0f} TiB ({usable_tb_floor:.0f} TB). "
            f"Nośniki NVMe muszą posiadać wbudowany mechanizm wykrywania zagrożeń ransomware "
            f"poprzez monitorowanie operacji zapisu — bez wpływu na wydajność."
        )
    else:
        _cap_txt = (
            f"Całkowita pojemność fizyczna (raw): min. {raw_tib_floor:.0f} TiB ({raw_tb_floor:.0f} TB). "
            f"Pojemność użyteczna (usable) w konfiguracji {raid_type}: "
            f"min. {usable_tib_floor:.0f} TiB ({usable_tb_floor:.0f} TB). "
            f"Wymagana minimalna liczba nośników NVMe: {drives_count} szt. "
            f"Nośniki NVMe muszą posiadać wbudowany mechanizm sprzętowego wykrywania "
            f"zagrożeń ransomware poprzez monitorowanie operacji zapisu — bez wpływu na wydajność."
        )

    requirements = [
        (
            1,
            "Obudowa",
            _enc_txt,
        ),
        (
            2,
            "Architektura",
            f"System musi składać się z pojedynczej macierzy dyskowej zarządzanej z jednego "
            f"interfejsu graficznego (GUI) oraz tekstowego (CLI). "
            + (
                f"Komunikacja kontrolerów z nośnikami NVMe odbywa się przez protokół NVMe; "
                f"nośniki HDD NL-SAS podłączone są przez dedykowaną kartę SAS. "
                f"Magistrala wewnętrzna NVMe: minimum PCIe 4.0."
                if (is_hybrid and hdd_drives_count > 0) else
                f"Komunikacja kontrolerów z nośnikami danych musi odbywać się wyłącznie przez protokół NVMe. "
                f"Zamawiający nie dopuszcza protokołu SAS do komunikacji wewnętrznej z nośnikami. "
                f"Magistrala wewnętrzna: minimum PCIe 4.0."
            ),
        ),
        (
            3,
            "Pojemność",
            _cap_txt,
        ),
        (
            4,
            "Kontrolery macierzowe",
            f"System musi być zbudowany z minimum dwóch kontrolerów pracujących w trybie active-active "
            f"lub dual-active bez pojedynczego punktu awarii (no SPOF). "
            f"Każdy kontroler musi być wyposażony w procesor klasy serwerowej; "
            f"system musi zapewniać łącznie min. {_cores_total} rdzeni CPU "
            f"we wszystkich kontrolerach. "
            f"Architektura procesorów: x86/x64.",
        ),
        (
            5,
            "Pamięć cache",
            f"System musi być wyposażony w minimum {cache_gb} GB pamięci podręcznej cache "
            f"(min. {cache_gb // 2} GB na kontroler). "
            f"Pamięć cache musi być zabezpieczona kondensatorem lub podtrzymaniem bateryjnym "
            f"chroniącym dane przy zaniku zasilania. "
            f"Elementy te muszą być wymienialne bez wyłączania systemu (hot-swap).",
        ),
        (
            6,
            "Interfejsy i protokoły",
            f"System w chwili dostawy musi posiadać minimum {fc_ports} portów FC 32 Gb/s "
            f"z możliwością rozbudowy o co najmniej kolejne 8 portów FC 32 Gb/s. "
            f"Porty FC muszą obsługiwać protokół FC-NVMe. "
            f"Wymagane interfejsy Ethernet 10/25 Gb/s obsługujące iSCSI oraz NVMe-oF. "
            f"Wraz z systemem należy dostarczyć {cable_label}.",
        ),
        (
            7,
            "Bezpieczeństwo danych",
            f"System musi obsługiwać poziomy RAID 1 oraz RAID 6 "
            f"z dystrybuowaną przestrzenią zapasową (bez dedykowanego fizycznego dysku spare — "
            f"przestrzeń spare rozłożona równomiernie na wszystkich napędach w puli) "
            f"i odpornością na jednoczesną awarię 2 napędów "
            f"(min. {rebuild_areas} obszary odbudowy). "
            f"Szyfrowanie danych: {enc_txt}. "
            f"Wymagana walidacja i szyfrowanie dysków systemowych kontrolerów (boot drive). "
            f"Obsługa dysków NVMe SED. Szybkość odbudowy RAID: min. 2 TB/h.",
        ),
        (
            8,
            "Niezawodność",
            f"Wszystkie krytyczne komponenty (adaptery HBA, kontrolery, pamięć, zasilacze, wentylatory) "
            f"muszą być nadmiarowe i wymienialne w trakcie pracy systemu (hot-swap). "
            f"Zasilanie z dwóch niezależnych źródeł jednofazowych 200–240 V / 50–60 Hz. "
            f"Wymagane baterie podtrzymujące zawartość pamięci cache przy zaniku zasilania "
            f"dla co najmniej dwóch kolejnych awarii zasilania.",
        ),
        (
            9,
            "Zarządzanie",
            f"Zarządzanie przez redundantne interfejsy Ethernet (min. 1 Gbps) z przeglądarki (HTTPS). "
            f"Obsługa zarządzania: portami We/Wy, woluminami, nośnikami NVMe, klonowaniem, "
            f"replikacją, wirtualizacją zasobów zewnętrznych. "
            f"Wymagana automatyzacja przez Ansible (oficjalny moduł producenta). "
            f"Obsługa wieloskładnikowej autentykacji (MFA). "
            f"Synchronizacja czasu NTP. Automatyczne raporty serwisowe (SMTP/HTTPS). "
            f"Powiadomienia SNMP.",
        ),
        (
            10,
            "Funkcjonalności",
            f"Wirtualizacja zasobów zewnętrznych (heterogeniczne macierze różnych producentów). "
            f"Klaster geograficzny (aktywno-aktywny) z replikacją do co najmniej trzeciego ośrodka. "
            f"Thin provisioning, kompresja inline, deduplikacja — wszystkie techniki równocześnie. "
            f"Dynamiczne zwiększanie rozmiaru wolumenów do min. 64 TB. "
            f"QoS na poziomie woluminu (IOPS oraz MB/s). "
            f"Wsparcie systemów operacyjnych: Windows Server 2016/2019/2022, RHEL 7.x/8.x, "
            f"SUSE Linux 12/15, VMware vSphere 7.0/8.0 i nowsze.",
        ),
        (
            11,
            "Wirtualne dyski logiczne",
            f"System musi obsługiwać woluminy rozłożone między co najmniej dwa różne typy "
            f"wirtualizowanych macierzy. Kopia lustrzana (mirror) woluminu między macierzami. "
            f"Obsługa min. 2 000 hostów i 64 000 mapowań wolumen-host.",
        ),
        (
            12,
            "Replikacja, klony, snapshoty",
            f"System musi obsługiwać replikację sterowaną politykami w trybie synchronicznym "
            f"(RPO = 0, active-active z automatycznym przełączeniem awaryjnym) "
            f"oraz asynchronicznym. "
            f"Wymagane spójne migawki grup woluminów (min. 256 na grupę), "
            f"kopie niezmienne (immutable) z retencją oraz klony binarne (incremental/multitarget). "
            f"Min. 2 500 relacji replikacji i 512 grup spójności.",
        ),
        (
            13,
            "Wydajność",
            perf_txt,
        ),
        (
            14,
            "Monitoring",
            f"Oprogramowanie do monitorowania musi przechowywać dane historyczne przez min. 365 dni. "
            f"Odpytywanie telemetrii w sekwencjach co najmniej 1-minutowych. "
            f"Monitorowanie metryk: IOPS, MB/s, czas odpowiedzi, cache, woluminy, hosty, przełączniki SAN. "
            f"Obsługa heterogenicznego środowiska (macierze różnych producentów). "
            f"Generowanie raportów chargeback. Integracja SNMP.",
        ),
        (
            15,
            "Inne wymagania",
            f"Macierz musi być fabrycznie nowa, nigdy wcześniej nie używana, "
            f"pochodzić z autoryzowanego kanału dystrybucji producenta "
            f"i być objęta serwisem producenta na terenie Rzeczpospolitej Polskiej. "
            f"Wraz z macierzą należy dostarczyć wszystkie elementy niezbędne do uruchomienia, "
            f"w tym wymagane licencje oprogramowania.",
        ),
        (
            16,
            "Serwis",
            f"Serwis świadczony przez producenta macierzy w trybie {sup_coverage}, {sup_fix}, "
            f"przez {sup_years} lat od daty dostawy. "
            f"Kontakt z pracownikiem serwisu w języku polskim przez 24 h na dobę, 7 dni w tygodniu. "
            f"Czas naprawy usterki krytycznej zgodny z fix-time SLA. "
            f"Uszkodzone nośniki danych stanowią własność Zamawiającego i nie podlegają zwrotowi.",
        ),
        (
            17,
            "Wbudowana inteligencja AI",
            f"System musi posiadać wbudowanego asystenta opartego na sztucznej inteligencji, "
            f"który wspomaga administratora w codziennych zadaniach operacyjnych: "
            f"diagnozowaniu problemów z wydajnością, rekomendacjach dotyczących konfiguracji, "
            f"analizie logów zdarzeń i prognozowaniu potencjalnych awarii. "
            f"Asystent AI musi działać w ramach oprogramowania zarządzającego macierzą, "
            f"bez konieczności instalacji dodatkowych komponentów. "
            f"Wymagana jest możliwość zadawania pytań w języku naturalnym dotyczących stanu systemu "
            f"oraz automatycznego generowania zaleceń konfiguracyjnych.",
        ),
        (
            18,
            "Gwarancja",
            f"Wymagana gwarancja producenta na wszystkie elementy systemu "
            f"(sprzęt i oprogramowanie) na okres {sup_years * 12} miesięcy ({sup_years} lat) "
            f"od daty podpisania protokołu odbioru.",
        ),
    ]

    return requirements


def _build_requirements_en(project: dict, model_info: dict, iops: int) -> list[tuple]:
    """English version of requirements — same structure, translated strings."""

    is_hybrid     = project.get("is_hybrid", False)
    drives_count  = project.get("drives_count", 0)
    raw_tib       = project.get("raw_tib",   0.0)
    raw_tb        = project.get("raw_tb",    0.0)
    usable_tib    = project.get("usable_tib",  0.0)
    usable_tb     = project.get("usable_tb",  0.0)
    cache_gb      = project.get("cache_gb",   256)
    fc_ports      = project.get("fc_ports",   8)
    raid_type     = project.get("raid_type",  "RAID 6")
    rebuild_areas = project.get("rebuild_areas", 2)
    form_factor   = model_info.get("form_factor", "1U")
    rack_units    = project.get("rack_units", None)
    encryption    = project.get("encryption", False)
    cable_qty     = project.get("cable_qty", 0)
    cable_label   = f"{cable_qty} FC cables" if cable_qty else "FC cables matching the number of ports"

    # Hybrid tier fields
    nvme_drives_count  = project.get("nvme_drives_count",  drives_count)
    nvme_raw_tib       = project.get("nvme_raw_tib",       0.0)
    nvme_raw_tb        = project.get("nvme_raw_tb",        0.0)
    nvme_usable_tib    = project.get("nvme_usable_tib",    0.0)
    nvme_usable_tb     = project.get("nvme_usable_tb",     0.0)
    drive_type         = project.get("drive_type",         "")
    hdd_drives_count   = project.get("hdd_drives_count",   0)
    hdd_raw_tib        = project.get("hdd_raw_tib",        0.0)
    hdd_raw_tb         = project.get("hdd_raw_tb",         0.0)
    hdd_usable_tib     = project.get("hdd_usable_tib",     0.0)
    hdd_usable_tb      = project.get("hdd_usable_tb",      0.0)
    hdd_drive_type     = project.get("hdd_drive_type",     "")
    hdd_enclosure      = project.get("hdd_enclosure",      "")
    hdd_enclosure_qty  = project.get("hdd_enclosure_qty",  1)

    model_short  = model_info.get("short", "")
    max_slots    = _MAX_DRIVE_SLOTS.get(model_short, max(drives_count, 12))
    _cores_total, _cores_desc = _CONTROLLER_CORES.get(model_short, (24, "min. 24 cores per controller pair"))

    raw_tib_floor    = _floor_tib(raw_tib)    or round(raw_tib,    1)
    raw_tb_floor     = _floor_tib(raw_tb)     or round(raw_tb,     1)
    usable_tib_floor = _floor_tib(usable_tib) or round(usable_tib, 1)
    usable_tb_floor  = _floor_tib(usable_tb)  or round(usable_tb,  1)

    enc_txt = (
        "Required — at-rest encryption without performance impact; "
        "support for external key servers (KMIP) and local USB keys"
        if encryption else
        "Required — encryption licence must be included; KMIP server support required"
    )

    sup          = project.get("support_info") or {}
    sup_years    = sup.get("years",    3)
    sup_coverage = sup.get("coverage", "24×7")
    sup_fix      = "with fix-time SLA" if sup.get("fix_time") else "without fix-time SLA"

    perf_iops_sub1 = project.get("perf_iops_max_sub1ms", 0)
    throughput_mib = project.get("perf_throughput_mib", 0.0)

    # Round down to nearest 25 000 IOPS and 100 MiB/s to avoid pinpointing exact vendor values
    iops_sub1_floor = (int(perf_iops_sub1 / 25_000) * 25_000) if perf_iops_sub1 else 0
    tp_floor        = (int(throughput_mib / 100) * 100)        if throughput_mib  else 0

    if iops_sub1_floor:
        perf_txt = (
            f"The system must deliver at least {iops_sub1_floor:,} IOPS "
            f"at sub-1 ms latency (measured for the configured workload profile)."
        )
        if tp_floor:
            perf_txt += f" Required sequential throughput: min. {tp_floor:,} MiB/s."
    elif iops:
        perf_txt = (
            f"The system must deliver a minimum of {iops:,} IOPS for the configured workload profile "
            f"(4 KiB block size, mixed read/write)."
        )
        if tp_floor:
            perf_txt += f" Sequential throughput: min. {tp_floor:,} MiB/s."
    else:
        perf_txt = (
            "The system must deliver performance appropriate to the Purchaser's requirements. "
            "Detailed IOPS requirements shall be completed based on a workload analysis."
        )

    # --- Enclosure text (EN) ---
    if is_hybrid and hdd_enclosure:
        _enc_txt = (
            f"The array must be installable in a standard 19\" rack. "
            f"The solution consists of a controller enclosure ({form_factor}) "
            f"and {hdd_enclosure_qty} × {hdd_enclosure} expansion enclosure(s) "
            f"housing NL-SAS HDD drives. "
            f"Total rack space: {rack_units} U. "
            f"The controller enclosure must contain at least two controllers "
            f"operating in active-active mode."
        )
    else:
        _enc_txt = (
            f"The array must be installable in a standard 19\" rack. "
            f"The base enclosure, no more than {form_factor} high, must contain "
            f"at least two controllers operating in active-active mode "
            f"and {max_slots} drive slots."
        )

    # --- Capacity text (EN) ---
    if is_hybrid and hdd_drives_count > 0:
        _nvme_raw_tib_fl  = _floor_tib(nvme_raw_tib)    or round(nvme_raw_tib,  1)
        _nvme_raw_tb_fl   = _floor_tib(nvme_raw_tb)     or round(nvme_raw_tb,   1)
        _nvme_us_tib_fl   = _floor_tib(nvme_usable_tib) or round(nvme_usable_tib, 1)
        _nvme_us_tb_fl    = _floor_tib(nvme_usable_tb)  or round(nvme_usable_tb,  1)
        _hdd_raw_tib_fl   = _floor_tib(hdd_raw_tib)     or round(hdd_raw_tib,   1)
        _hdd_raw_tb_fl    = _floor_tib(hdd_raw_tb)      or round(hdd_raw_tb,    1)
        _hdd_us_tib_fl    = _floor_tib(hdd_usable_tib)  or round(hdd_usable_tib, 1)
        _hdd_us_tb_fl     = _floor_tib(hdd_usable_tb)   or round(hdd_usable_tb,  1)
        _enc_str          = f" in {hdd_enclosure} enclosure" if hdd_enclosure else ""
        _cap_txt = (
            f"The system must feature a two-tier capacity architecture (NVMe Tier + HDD Tier). "
            f"NVMe Tier: {nvme_drives_count} NVMe drives, "
            f"raw capacity min. {_nvme_raw_tib_fl:.0f} TiB ({_nvme_raw_tb_fl:.0f} TB), "
            f"usable capacity min. {_nvme_us_tib_fl:.0f} TiB ({_nvme_us_tb_fl:.0f} TB). "
            f"HDD Tier: {hdd_drives_count} NL-SAS drives{_enc_str}, "
            f"raw capacity min. {_hdd_raw_tib_fl:.0f} TiB ({_hdd_raw_tb_fl:.0f} TB), "
            f"usable capacity min. {_hdd_us_tib_fl:.0f} TiB ({_hdd_us_tb_fl:.0f} TB). "
            f"Total raw capacity: min. {raw_tib_floor:.0f} TiB ({raw_tb_floor:.0f} TB). "
            f"Total usable capacity ({raid_type}): min. {usable_tib_floor:.0f} TiB ({usable_tb_floor:.0f} TB). "
            f"NVMe drives must include a built-in ransomware detection mechanism "
            f"by monitoring write operations — without affecting performance."
        )
    else:
        _cap_txt = (
            f"Total physical (raw) capacity: min. {raw_tib_floor:.0f} TiB ({raw_tb_floor:.0f} TB). "
            f"Usable capacity in {raid_type} configuration: "
            f"min. {usable_tib_floor:.0f} TiB ({usable_tb_floor:.0f} TB). "
            f"Minimum number of NVMe drives required: {drives_count}. "
            f"NVMe drives must include a built-in hardware-accelerated ransomware detection mechanism "
            f"by monitoring write operations — without affecting performance."
        )

    requirements = [
        (1,  "Enclosure", _enc_txt),
        (2,  "Architecture",
         f"The system must consist of a single storage array managed from a single "
         f"graphical (GUI) and command-line (CLI) interface. "
         + (
             f"Controller-to-NVMe-drive communication uses the NVMe protocol; "
             f"NL-SAS HDD drives are connected via a dedicated SAS adapter. "
             f"Internal NVMe bus: minimum PCIe 4.0."
             if (is_hybrid and hdd_drives_count > 0) else
             f"Controller-to-drive communication must use the NVMe protocol exclusively. "
             f"SAS for internal drive communication is not permitted. "
             f"Internal bus: minimum PCIe 4.0."
         )),
        (3,  "Capacity", _cap_txt),
        (4,  "Storage Controllers",
         f"The system must include at least two controllers operating in active-active "
         f"or dual-active mode with no single point of failure (no SPOF). "
         f"Each controller must be equipped with a server-class processor; "
         f"the system must provide a minimum of {_cores_total} CPU cores in total "
         f"across all controllers. "
         f"Processor architecture: x86/x64."),
        (5,  "Cache Memory",
         f"The system must have at least {cache_gb} GB of cache memory "
         f"(min. {cache_gb // 2} GB per controller). "
         f"Cache must be protected by a capacitor or battery backup preserving data "
         f"on power loss. These components must be hot-swappable."),
        (6,  "Interfaces & Protocols",
         f"At delivery the system must have at least {fc_ports} × 32 Gb/s FC ports "
         f"with the ability to add at least 8 more 32 Gb/s FC ports. "
         f"FC ports must support the FC-NVMe protocol. "
         f"Ethernet 10/25 Gb/s interfaces supporting iSCSI and NVMe-oF are required. "
         f"The following must be included with the system: {cable_label}."),
        (7,  "Data Security",
         f"The system must support RAID 1 and RAID 6 "
         f"with distributed spare capacity (no dedicated physical spare drive — "
         f"spare space distributed evenly across all drives in the pool) "
         f"and tolerance for simultaneous failure of 2 drives "
         f"(min. {rebuild_areas} rebuild areas). "
         f"Data encryption: {enc_txt}. "
         f"Validation and encryption of controller system drives (boot drive) is required. "
         f"Support for NVMe SED drives. RAID rebuild speed: min. 2 TB/h."),
        (8,  "Reliability",
         f"All critical components (HBA adapters, controllers, memory, PSUs, fans) "
         f"must be redundant and hot-swappable. "
         f"Dual independent single-phase power feeds 200–240 V / 50–60 Hz. "
         f"Battery backup protecting cache content on power loss "
         f"for at least two consecutive power failures."),
        (9,  "Management",
         f"Management via redundant Ethernet interfaces (min. 1 Gbps) from a browser (HTTPS). "
         f"Management must cover: I/O ports, volumes, NVMe drives, cloning, "
         f"replication, external storage virtualisation. "
         f"Automation via Ansible (official vendor module) is required. "
         f"Multi-factor authentication (MFA) support. "
         f"NTP time synchronisation. Automated service reports (SMTP/HTTPS). "
         f"SNMP notifications."),
        (10, "Functionality",
         f"External storage virtualisation (heterogeneous arrays from multiple vendors). "
         f"Geographic cluster (active-active) with replication to at least a third site. "
         f"Thin provisioning, inline compression, deduplication — all simultaneously. "
         f"Dynamic volume resizing up to min. 64 TB. "
         f"QoS at volume level (IOPS and MB/s). "
         f"OS support: Windows Server 2016/2019/2022, RHEL 7.x/8.x, "
         f"SUSE Linux 12/15, VMware vSphere 7.0/8.0 and later."),
        (11, "Logical Volumes",
         f"The system must support volumes spanning at least two different "
         f"virtualised array types. Mirror copy of a volume between arrays. "
         f"Support for min. 2,000 hosts and 64,000 volume-to-host mappings."),
        (12, "Replication, Clones & Snapshots",
         f"The system must support policy-driven replication in synchronous mode "
         f"(RPO = 0, active-active with automatic failover) and asynchronous mode. "
         f"Consistent snapshots of volume groups are required (min. 256 per group), "
         f"immutable copies with retention, and binary clones (incremental / multi-target). "
         f"Min. 2,500 replication relationships and 512 consistency groups."),
        (13, "Performance",
         perf_txt),
        (14, "Monitoring",
         f"Monitoring software must retain historical data for at least 365 days. "
         f"Telemetry polling at minimum 1-minute intervals. "
         f"Metrics: IOPS, MB/s, response time, cache, volumes, hosts, SAN switches. "
         f"Heterogeneous environment support (multiple vendors). "
         f"Chargeback report generation. SNMP integration."),
        (15, "Other Requirements",
         f"The array must be brand new, never previously used, "
         f"sourced through the vendor's authorised distribution channel, "
         f"and covered by the vendor's support within the territory. "
         f"All components required for commissioning must be delivered with the array, "
         f"including required software licences."),
        (16, "Service",
         f"Service provided by the array manufacturer in {sup_coverage} mode, {sup_fix}, "
         f"for {sup_years} years from the delivery date. "
         f"Contact with a service engineer in English available 24 hours a day, 7 days a week. "
         f"Critical fault repair time in accordance with fix-time SLA. "
         f"Failed storage media remain the property of the Purchaser and are not to be returned."),
        (17, "Built-in AI Intelligence",
         f"The system must include a built-in artificial intelligence assistant "
         f"that helps the administrator with daily operational tasks: "
         f"diagnosing performance issues, providing configuration recommendations, "
         f"analysing event logs and predicting potential failures. "
         f"The AI assistant must operate within the array management software "
         f"without requiring additional component installation. "
         f"The ability to ask natural-language questions about system status "
         f"and to automatically generate configuration recommendations is required."),
        (18, "Warranty",
         f"A manufacturer warranty on all system components "
         f"(hardware and software) for {sup_years * 12} months ({sup_years} years) "
         f"from the date of acceptance protocol sign-off is required."),
    ]

    return requirements


# ---------------------------------------------------------------------------
# DOCX helpers (mirrors exec_summary.py)
# ---------------------------------------------------------------------------

def _set_page_margins(doc):
    from docx.oxml.ns import qn as _qn
    section = doc.sections[0]
    section.page_width   = Cm(21)
    section.page_height  = Cm(29.7)
    section.left_margin  = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin   = Cm(2.0)
    section.bottom_margin= Cm(2.0)


def _set_default_font(doc):
    from docx.oxml.ns import qn as _qn
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _para_space(p, before=0, after=6):
    from docx.oxml.ns import qn as _qn
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(_qn("w:before"), str(before * 20))
    spacing.set(_qn("w:after"),  str(after  * 20))
    pPr.append(spacing)


def _para_space_after_table(doc):
    p = doc.add_paragraph()
    _para_space(p, before=0, after=4)


def _add_hrule(doc):
    p = doc.add_paragraph()
    _para_space(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0062FF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(
    cell, text: str,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    size: int = 9,
    bold: bool = False,
    color: RGBColor | None = None,
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    _para_space(p, before=2, after=2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = IBM_DARK
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
