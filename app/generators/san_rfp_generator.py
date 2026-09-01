"""
SAN Switch RFP / Technical Requirements Generator.
Produces a vendor-neutral Fibre Channel SAN fabric specification document.
All references to OEM origin (Brocade/Broadcom) are omitted — IBM b-type switches
are listed as IBM products.
"""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

IBM_BLUE  = RGBColor(0x00, 0x62, 0xFF)
IBM_DARK  = RGBColor(0x16, 0x16, 0x16)
IBM_GRAY  = RGBColor(0x52, 0x52, 0x52)
IBM_LG    = RGBColor(0xF4, 0xF4, 0xF4)
IBM_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_san_rfp(
    project: dict[str, Any],
    client_name: str = "",
    seller_name: str = "",
    lang: str = "en",
) -> bytes:
    """
    Generate a short SAN infrastructure requirements document.
    Works for SAN-only and is also called to append a SAN section to FS RFPs.
    Returns DOCX bytes.
    """
    doc = Document()
    _set_page_margins(doc)
    _set_default_font(doc)

    san_switches = project.get("san_switches", [])

    _add_header_block(doc, client_name, lang)
    _add_intro(doc, project, san_switches, client_name, lang)
    _add_requirements_table(doc, san_switches, lang)
    _add_support_section(doc, san_switches, lang)
    _add_footer(doc, lang)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _add_header_block(doc, client_name: str, lang: str) -> None:
    if lang == "pl":
        title_txt  = "Specyfikacja Techniczna — Wymagania dla Sieci SAN"
        client_lbl = "Zamawiający"
        date_lbl   = "Data"
        date_fmt   = "%d.%m.%Y"
    else:
        title_txt  = "Technical Specification — Fibre Channel SAN Fabric Requirements"
        client_lbl = "Client"
        date_lbl   = "Date"
        date_fmt   = "%B %d, %Y"

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _para_space(title, before=6, after=2)
    r = title.add_run(title_txt)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = IBM_DARK

    sub = doc.add_paragraph()
    _para_space(sub, before=0, after=2)
    r = sub.add_run(
        f"{client_lbl}: {client_name or '—'}  ·  "
        f"{date_lbl}: {date.today().strftime(date_fmt)}"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = IBM_GRAY
    _add_hrule(doc)


def _add_intro(doc, project: dict, san_switches: list, client_name: str, lang: str) -> None:
    if lang == "pl":
        sup = (san_switches[0].get("support_info") or {}) if san_switches else {}
        sup_name = sup.get("name", "IBM Storage Expert Care")
        sup_yrs  = sup.get("years", "5")
        total_qty = sum(sw.get("qty", 1) for sw in san_switches)
        txt = (
            f"Niniejszy dokument określa minimalne wymagania techniczne dla infrastruktury "
            f"sieci Fibre Channel SAN dla {client_name or 'zamawiającego'}. "
            f"Zamówienie obejmuje dostawę {total_qty} przełącznika/przełączników FC "
            f"wraz z oprogramowaniem zarządzającym, okablowaniem i wsparciem technicznym "
            f"{sup_name} na okres {sup_yrs} lat."
        )
    else:
        sup = (san_switches[0].get("support_info") or {}) if san_switches else {}
        sup_name = sup.get("name", "IBM Storage Expert Care")
        sup_yrs  = sup.get("years", "5")
        total_qty = sum(sw.get("qty", 1) for sw in san_switches)
        txt = (
            f"This document defines minimum technical requirements for the Fibre Channel "
            f"SAN fabric infrastructure for {client_name or 'the client'}. "
            f"The procurement covers the supply of {total_qty} FC switch(es) "
            f"including management software, cabling and {sup_name} hardware support "
            f"for {sup_yrs} years."
        )
    p = doc.add_paragraph()
    _para_space(p, before=4, after=6)
    r = p.add_run(txt)
    r.font.size = Pt(10)
    r.font.color.rgb = IBM_DARK


def _add_requirements_table(doc, san_switches: list, lang: str) -> None:
    # Table heading — matches FS RFP style (Pt(11), bold, IBM_DARK, before=8, after=4)
    heading = doc.add_paragraph()
    _para_space(heading, before=8, after=4)
    run = heading.add_run(
        "Tabela Wymagań Technicznych" if lang == "pl" else "Technical Requirements Table"
    )
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = IBM_DARK

    if lang == "pl":
        h0, h1, h2 = "Lp.", "Parametr", "Minimalna specyfikacja"
    else:
        h0, h1, h2 = "No.", "Parameter / Element", "Minimum Specification"

    rows: list[tuple[str, str]] = []
    if san_switches:
        sw = san_switches[0]
        qty   = sum(s.get("qty", 1) for s in san_switches)
        speed = sw.get("port_speed_gbps", 32)
        max_p = sw.get("max_ports", 0)
        act_p = sw.get("active_ports", 0)
        form  = sw.get("form_factor", "1U")
        lw    = sum(s.get("lw_optics_qty", 0) for s in san_switches)
        sw_c  = sum(s.get("sw_optics_qty", 0) for s in san_switches)

        if lang == "pl":
            rows = [
                ("Ilość przełączników",   f"min. {qty} szt."),
                ("Protokół",              "Fibre Channel (FC) Gen 7 lub nowszy"),
                ("Prędkość portów",       f"min. {speed} Gbps FC auto-negotiation"),
                ("Porty aktywne",         f"min. {act_p} portów FC (maks. {max_p})"),
                ("Format obudowy",        f"{form} rack-mount, zasilanie redundantne hot-swap"),
                ("Optyki LW",             f"{lw} szt. SFP+ jednomodowy ≥10 km" if lw else "Nie dotyczy"),
                ("Kable SW / OM3",        f"{sw_c} szt. kabel OM3 LC/LC" if sw_c else "Nie dotyczy"),
                ("Zarządzanie",           "GUI; CLI; SNMP v3; syslog; REST API"),
                ("NVMe-oF",               "Obsługa FC-NVMe (NVMe over Fibre Channel)"),
                ("Wirtualizacja sieci",   "NPIV; trunking ISL; obsługa FICON (opcja)"),
                ("Szyfrowanie ISL",       "AES-256 na łączach ISL (opcja)"),
                ("Aktualizacje FW",       "Nieprzerywające pracy (non-disruptive firmware upgrade)"),
            ]
        else:
            rows = [
                ("Quantity",              f"min. {qty} unit(s)"),
                ("Protocol",              "Fibre Channel (FC) Gen 7 or later"),
                ("Port speed",            f"min. {speed} Gbps FC, auto-negotiation"),
                ("Active ports",          f"min. {act_p} FC ports licensed (expandable to {max_p})"),
                ("Form factor",           f"{form} rack-mount, hot-swap redundant PSUs"),
                ("LW optics (long-wave)", f"{lw} × SFP+ single-mode ≥10 km" if lw else "N/A"),
                ("SW cables / OM3",       f"{sw_c} × OM3 LC/LC multimode cable" if sw_c else "N/A"),
                ("Management",            "GUI; CLI; SNMP v3; syslog; REST API"),
                ("NVMe-oF readiness",     "FC-NVMe (NVMe over Fibre Channel) support required"),
                ("Fabric virtualisation", "NPIV; ISL trunking; FICON-capable (optional)"),
                ("ISL encryption",        "AES-256 on ISL links (optional)"),
                ("Firmware upgrade",      "Non-disruptive in-service firmware upgrade"),
            ]

    # 3-column table: No. | Parameter | Specification — same style as FS RFP
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Column widths — identical to FS RFP [Cm(1.2), Cm(4.5), Cm(12.8)]
    col_widths = [Cm(1.2), Cm(4.5), Cm(12.8)]

    # Header — dark background (matches FS RFP header style)
    hdr = table.add_row()
    for ci, w in enumerate(col_widths):
        hdr.cells[ci].width = w
    _set_cell_bg(hdr.cells[0], IBM_DARK)
    _set_cell_bg(hdr.cells[1], IBM_DARK)
    _set_cell_bg(hdr.cells[2], IBM_DARK)
    _set_cell_text(hdr.cells[0], h0,  bold=True, color=IBM_WHITE, size=9)
    _set_cell_text(hdr.cells[1], h1,  bold=True, color=IBM_WHITE, size=9)
    _set_cell_text(hdr.cells[2], h2,  bold=True, color=IBM_WHITE, size=9)

    # Data rows — zebra stripe
    for i, (label, val) in enumerate(rows):
        row = table.add_row()
        bg = IBM_LG if i % 2 == 0 else IBM_WHITE
        for ci, w in enumerate(col_widths):
            row.cells[ci].width = w
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        _set_cell_bg(row.cells[2], bg)
        _set_cell_text(row.cells[0], str(i + 1), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(row.cells[1], label, bold=True, size=9)
        _set_cell_text(row.cells[2], val, size=9)

    _para_space_after_table(doc)


def _add_support_section(doc, san_switches: list, lang: str) -> None:
    if not san_switches:
        return
    sup = (san_switches[0].get("support_info") or {})
    sup_name  = sup.get("name", "IBM Storage Expert Care")
    sup_yrs   = sup.get("years", "5")
    sup_cover = sup.get("coverage", "24×7")
    fix_h     = sup.get("fix_time_hours", "")
    fix_str   = f", {fix_h} on-site response" if fix_h else ""

    if lang == "pl":
        _section_heading(doc, "Wsparcie techniczne")
        txt = (
            f"Wymagane wsparcie techniczne {sup_name} na okres {sup_yrs} lat, "
            f"pokrycie {sup_cover}{fix_str}. "
            f"Wsparcie musi zapewniać proaktywny monitoring infrastruktury oraz "
            f"dostęp do aktualnych aktualizacji oprogramowania sprzętowego."
        )
    else:
        _section_heading(doc, "Support Requirements")
        txt = (
            f"Required hardware support: {sup_name}, {sup_yrs}-year term, "
            f"{sup_cover}{fix_str}. "
            f"Support must include proactive monitoring and firmware update entitlement."
        )
    p = doc.add_paragraph()
    _para_space(p, before=0, after=6)
    r = p.add_run(txt)
    r.font.size = Pt(10)
    r.font.color.rgb = IBM_DARK


def _add_footer(doc, lang: str) -> None:
    _add_hrule(doc)
    note = (
        "Wymagania techniczne opisane w niniejszym dokumencie mają charakter neutralny technologicznie. "
        "Zamawiający zastrzega sobie prawo weryfikacji zgodności oferowanego rozwiązania."
        if lang == "pl" else
        "All technical requirements in this document are technology-neutral. "
        "The client reserves the right to verify compliance of the proposed solution."
    )
    p = doc.add_paragraph()
    _para_space(p, before=4, after=0)
    r = p.add_run(note)
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = IBM_GRAY


# ---------------------------------------------------------------------------
# Helpers (shared formatting)
# ---------------------------------------------------------------------------

def _section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    _para_space(p, before=10, after=3)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = IBM_BLUE


def _set_page_margins(doc) -> None:
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.0)
    sec.right_margin  = Cm(2.0)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)


def _set_default_font(doc) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _para_space(p, before=0, after=6) -> None:
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"),  str(after  * 20))
    pPr.append(spacing)


def _para_space_after_table(doc) -> None:
    p = doc.add_paragraph()
    _para_space(p, before=0, after=4)


def _add_hrule(doc) -> None:
    p = doc.add_paragraph()
    _para_space(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0062FF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_bg(cell, color: RGBColor) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False,
                   color: RGBColor = IBM_DARK,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    _para_space(p, before=2, after=2)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = color
