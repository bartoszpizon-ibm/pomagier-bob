"""
Special Bid Pricing Request Questionnaire Generator.
Opens the original DOCX template and fills in the answer cells in-place,
preserving all formatting, layout and branding of the original document.
"""

from __future__ import annotations

import copy
import io
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.knowledge.product_db import get_model_info

TEMPLATE_PATH = Path(__file__).parent.parent.parent / \
    "HW Special Bid Pricing Request Questionnaire - TEST.docx"

# Colour used for the answer text — plain black, same as "Answer:" label
_BLACK = RGBColor(0x00, 0x00, 0x00)
_BLUE  = RGBColor(0x36, 0x5F, 0x91)   # same blue as best-practice examples


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_special_bid(
    project: dict[str, Any],
    client_name: str = "",
    seller_name: str = "",
    distributor_name: str = "",
    reseller_name: str = "",
    discount_pct: float = 60.0,
    opportunity_context: str = "",
    deal_background: str = "",
    competitor_info: str = "",
    deal_history: str = "",
    business_justification: str = "",
    extended_validity_days: int = 0,
    extended_validity_reason: str = "",
    num_systems: int = 1,
    eu_margin_pct: float = 15.0,
) -> bytes:
    """
    Fill the original Special Bid DOCX template and return bytes.

    Table map (0-based):
      Table 0 — 'Required field' banner  (skip)
      Table 1 — Tier / channel table     (fill col 1, rows 1-3 + row 5)
      Table 2 — Answer Q1 Opportunity    (row 0 = Answer, row 1 = example → replace)
      Table 3 — Answer Q2 Background     (row 0 = Answer, row 1 = example → replace)
      Table 4 — Pricing Summary          (col 1: rows 1-3 → fill values)
      Table 5 — Answer Q3 Pricing just.  (row 0 = Answer, row 1 = example → replace)
      Table 6 — Answer Q4 Competitors    (row 0 = Answer, row 1 = example → replace)
      Table 7 — Answer C Deal history    (row 0 = Answer, row 1 = example → replace)
    """
    doc = Document(str(TEMPLATE_PATH))

    model_code = project.get("model_code", "")
    model_info = get_model_info(model_code)
    n = max(1, int(num_systems))
    model_name = model_info.get("name", model_code)

    # --- build pricing strings (scaled by number of systems) ---
    list_hw  = project.get("list_price_hw", 0.0)
    list_sw  = project.get("list_price_sw", 0.0)
    list_sup = project.get("list_price_support", 0.0)
    ship     = project.get("shipping", 0.0)
    curr     = project.get("currency", "EUR")
    d        = discount_pct / 100
    m        = 1 + eu_margin_pct / 100
    net_hw   = list_hw  * (1 - d) * n
    net_sw   = list_sw  * (1 - d) * n
    net_sup  = list_sup * (1 - d) * n
    net_ship = ship * n
    net_tot  = net_hw + net_sw + net_sup + net_ship      # BP total
    eu_tot   = (net_hw + net_sw + net_sup) * m + net_ship  # End User total (with margin)
    list_tot = (list_hw + list_sw + list_sup + ship) * n

    # --- prepend quantity note to opportunity_context when n > 1 ---
    if n > 1:
        _qty_prefix = (
            f"This bid covers the supply of {n} × {model_name} units. "
            f"All pricing below reflects the total for {n} systems. "
        )
        opportunity_context = _qty_prefix + (opportunity_context or "")

    # --- auto-build pricing justification if blank ---
    if not business_justification:
        dev = discount_pct - 60.0
        dev_str = (f"a {dev:.1f}-point deviation above the standard 60% baseline"
                   if dev > 0 else "within the standard 60% baseline")
        _sys_str = f" for {n} × {model_name}" if n > 1 else ""
        business_justification = (
            f"Requested BP price: {net_tot:,.0f} {curr}{_sys_str} "
            f"(IBM list: {list_tot:,.0f} {curr}) — discount {discount_pct:.1f}%, {dev_str}.\n\n"
            f"Justification: IBM list pricing is not competitive for this opportunity without "
            f"exception support. Competing vendors are expected to submit proposals priced "
            f"significantly below IBM list; failing to match their price band will result in "
            f"IBM losing the deal. The requested discount is the minimum level required to "
            f"align the IBM net price with the customer's confirmed budget and the competitive "
            f"price range established through the RFP benchmarking process.\n\n"
            f"IBM {model_name} value justification: (1) Hardware-accelerated inline AI ransomware "
            f"detection at the drive level — no additional software cost; "
            f"(2) Distributed RAID 6 with >2 TB/h rebuild speed — minimising data "
            f"exposure during drive failure for Tier-1 workloads; (3) IBM Storage Insights — "
            f"proactive capacity and performance management included with the purchase.\n\n"
            f"[Pricing approver note: please add specific competitor price intelligence and "
            f"customer budget confirmation before submission.]"
        )

    # Bid validity — extend if needed
    if extended_validity_days > 30:
        validity_days_used = extended_validity_days
    else:
        validity_days_used = 30
    valid_until = (date.today() + timedelta(days=validity_days_used)).strftime("%d %B %Y")

    # If extended validity requested, append note to opportunity_context and business_justification
    if extended_validity_days > 30 and extended_validity_reason:
        _ext_note = (
            f"\n\nNote — Extended Bid Validity: This offer requires a validity period of "
            f"{extended_validity_days} days (valid until {valid_until}) instead of the "
            f"standard 30 days. Reason: {extended_validity_reason}."
        )
        opportunity_context = (opportunity_context or "") + _ext_note
        business_justification = (business_justification or "") + (
            f"\n\nExtended Validity Justification: {extended_validity_reason}. "
            f"Requested validity: {extended_validity_days} days (until {valid_until})."
        )

    # ── Table 1 — Channel / Tier ─────────────────────────────────────────
    t1 = doc.tables[1]
    _fill_tier_cell(t1.rows[1].cells[1], distributor_name or "—")   # Tier 1 Distributor
    _fill_tier_cell(t1.rows[2].cells[1], reseller_name   or "—")    # Tier 2 Reseller
    _fill_tier_cell(t1.rows[3].cells[1], client_name     or "—")    # End User
    _fill_tier_cell(t1.rows[5].cells[1], seller_name     or "—")    # IBM Sales Rep

    # ── Table 2 — Q1 Opportunity Context ────────────────────────────────
    _fill_answer_table(doc.tables[2], opportunity_context,
                       valid_until=valid_until)

    # ── Table 3 — Q2 Deal Background ────────────────────────────────────
    _fill_answer_table(doc.tables[3], deal_background)

    # ── Table 4 — Pricing Summary ────────────────────────────────────────
    t4 = doc.tables[4]
    qty_label = f"{n} × {model_name}" if n > 1 else model_name

    # Row 1 — End User Buy Price (BP price + EU margin, same as End User Price in Exec Summary)
    _fill_pricing_cell(t4.rows[1].cells[1],
                       f"{eu_tot:,.2f} {curr}  ({qty_label})")
    # Row 2 — Requested End User Discount %
    _fill_pricing_cell(t4.rows[2].cells[1],
                       f"{discount_pct:.1f}%")
    # Row 3 — Total Value Seller Discount (standard 60%)
    _fill_pricing_cell(t4.rows[3].cells[1],
                       "60.0%")

    # ── Table 5 — Q3 Business Justification ─────────────────────────────
    _fill_answer_table(doc.tables[5], business_justification)

    # ── Table 6 — Q4 Competitive Positioning ────────────────────────────
    _fill_answer_table(doc.tables[6], competitor_info or
                       "No competitive data provided — please complete manually.")

    # ── Table 7 — C Deal History ─────────────────────────────────────────
    _fill_answer_table(doc.tables[7], deal_history or
                       "No prior related bids or transactions identified for this opportunity.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Cell-filling helpers
# ---------------------------------------------------------------------------

def _clear_cell(cell) -> None:
    """Remove all paragraphs from a cell, leaving exactly one empty one."""
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)
    # Add back one clean paragraph
    new_p = OxmlElement("w:p")
    tc.append(new_p)


def _fill_tier_cell(cell, text: str) -> None:
    """Write a plain bold 10pt black text into a tier-table name cell."""
    _clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _BLACK


def _fill_pricing_cell(cell, text: str) -> None:
    """Write pricing value into input cell (replaces [$] / [%] placeholder)."""
    _clear_cell(cell)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _BLACK


def _fill_answer_table(table, answer_text: str, valid_until: str = "") -> None:
    """
    Replace the content of an Answer table:
      row 0 = 'Answer:' label  → keep as-is, append the answer text after it
      row 1 = 'Best practice example' → replace with filled answer

    Strategy: write the answer into row 0 (after the 'Answer:' label) AND
    clear row 1 (best-practice example) replacing it with a brief note.
    """
    if not table.rows:
        return

    # --- Row 0: find 'Answer:' paragraph, clear remaining runs, add answer ---
    cell0 = table.rows[0].cells[0]
    _replace_answer_cell(cell0, answer_text, valid_until)

    # --- Row 1: clear best-practice example row entirely ---
    if len(table.rows) > 1:
        cell1 = table.rows[1].cells[0]
        _clear_cell(cell1)


def _replace_answer_cell(cell, answer_text: str, valid_until: str = "") -> None:
    """
    Keep the 'Answer:' bold label in its original paragraph,
    then append the answer text as subsequent paragraphs in the same cell.
    """
    tc = cell._tc
    paragraphs = tc.findall(qn("w:p"))

    # Find the paragraph that contains 'Answer:' and keep only it
    answer_para_xml = None
    for p_el in paragraphs:
        full_text = "".join(
            r.text for r in p_el.findall(".//" + qn("w:t"))
        )
        if "Answer" in full_text:
            answer_para_xml = p_el
            break

    # Remove all paragraphs from cell
    for p_el in list(paragraphs):
        tc.remove(p_el)

    # Re-add the Answer: paragraph (or create one if not found)
    if answer_para_xml is not None:
        # strip all runs after "Answer:" so we start clean
        for r in list(answer_para_xml.findall(qn("w:r"))):
            answer_para_xml.remove(r)
        # rebuild: bold "Answer:" run
        r_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b"); rPr.append(b)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "20"); rPr.append(sz)
        szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20"); rPr.append(szCs)
        r_el.append(rPr)
        t_el = OxmlElement("w:t"); t_el.text = "Answer:"; r_el.append(t_el)
        answer_para_xml.append(r_el)
        tc.append(answer_para_xml)
    else:
        new_p = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b"); rPr.append(b)
        r_el.append(rPr)
        t_el = OxmlElement("w:t"); t_el.text = "Answer:"; r_el.append(t_el)
        new_p.append(r_el)
        tc.append(new_p)

    # Add answer text as separate paragraphs (split on \n)
    lines = (answer_text or "").split("\n")
    for line in lines:
        p_new = OxmlElement("w:p")
        if line.strip():
            r_new = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            sz = OxmlElement("w:sz");   sz.set(qn("w:val"), "20"); rPr.append(sz)
            szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "20"); rPr.append(szCs)
            r_new.append(rPr)
            t_new = OxmlElement("w:t")
            t_new.text = line
            t_new.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r_new.append(t_new)
            p_new.append(r_new)
        tc.append(p_new)

    # Optionally append valid-until line
    if valid_until:
        p_v = OxmlElement("w:p")
        r_v = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        i = OxmlElement("w:i"); rPr.append(i)
        sz = OxmlElement("w:sz");   sz.set(qn("w:val"), "18"); rPr.append(sz)
        r_v.append(rPr)
        t_v = OxmlElement("w:t")
        t_v.text = f"Offer valid until: {valid_until}"
        r_v.append(t_v)
        p_v.append(r_v)
        tc.append(p_v)
