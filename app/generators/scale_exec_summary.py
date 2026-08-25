"""
Executive Summary DOCX generator for IBM Storage Scale System (3500 / 6000).
Produces an IBM-branded Word document from a parsed project dict
returned by app.parsers.scale_parser.parse_scale_project().
"""
from __future__ import annotations

import io
import re
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches

from app.knowledge.product_db import get_model_info, get_docs

# IBM design tokens
IBM_BLUE       = RGBColor(0x00, 0x62, 0xFF)
IBM_DARK       = RGBColor(0x16, 0x16, 0x16)
IBM_GRAY       = RGBColor(0x52, 0x52, 0x52)
IBM_LIGHT_GRAY = RGBColor(0xF4, 0xF4, 0xF4)
IBM_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

ASSETS_DIR = Path(__file__).parent.parent / "assets"
LOGOS_DIR  = ASSETS_DIR / "logos"
IMAGES_DIR = ASSETS_DIR / "images"

# ---------------------------------------------------------------------------
# Translations
# ---------------------------------------------------------------------------
_TRANS: dict[str, dict[str, str]] = {
    "en": {
        "cover_subtitle":    "Technical Executive Summary",
        "prepared_for":      "Prepared for",
        "prepared_by":       "Prepared by",
        "date":              "Date",
        "valid_until":       "Valid until",
        "config_id":         "Configuration ID",
        "sec_exec":          "Executive Summary",
        "sec_config":        "Solution Configuration",
        "sec_capacity":      "Capacity",
        "sec_performance":   "Performance Profile",
        "sec_connectivity":  "Connectivity & Data Protection",
        "sec_environment":   "Physical Environment",
        "sec_software":      "Software & Subscriptions",
        "sec_support":       "Service & Support",
        "sec_pricing":       "Pricing Summary",
        "sec_next":          "Next Steps",
        "key_highlights":    "Key Solution Highlights",
        "platform_advantages": "Platform Advantages",
        "adv1": (
            "Parallel file system with global namespace — all compute nodes access "
            "the same high-performance namespace simultaneously, eliminating I/O bottlenecks "
            "in AI, HPC, and analytics pipelines."
        ),
        "adv2": (
            "Erasure-coded data protection — distributes parity across all drives in the cluster "
            "without dedicating physical spare drives, maximising usable capacity efficiency."
        ),
        "adv3": (
            "Active-active scale-out architecture — add storage nodes non-disruptively; "
            "performance and capacity scale linearly while clients continue I/O without interruption."
        ),
        "adv4": (
            "Transparent cloud tiering — inactive data migrates automatically to object storage "
            "(on-premises or cloud), reducing primary storage footprint and cost without "
            "changing the access path for applications."
        ),
        "body": (
            "This proposal {client_str}recommends the {model_name} parallel file storage solution. "
            "The configuration comprises {num_nodes} storage node(s) with {drives_count} NVMe drives "
            "({drive_type}), delivering {usable_tib:.1f} TiB usable capacity. "
            "All hardware is covered by {support_name} providing {support_hours} for {support_years} years."
        ),
        "multi_system_note": "This document covers a configuration of {n} × {model}. Specifications below refer to a single system.",
        "support_24x7":      "24×7 around-the-clock support with a hardware fix-time SLA",
        "support_9x5":       "9×5 business-hours support",
        "no_perf_data":      "Performance data not provided. Upload a Storage Modeller performance report or enter throughput manually.",
        "no_support":        "No support information found in configuration.",
        "sup_pkg":           "Support Package",
        "sup_level":         "Level",
        "sup_term":          "Term",
        "sup_years_unit":    "{n} years",
        "sup_coverage":      "Coverage Hours",
        "sup_fixtime":       "Hardware Fix-Time SLA",
        "sup_fixtime_yes":   "Yes — next business day or same day (by agreement)",
        "sup_fixtime_no":    "No fix-time SLA",
        "sup_desc":          "Description",
        "sup_comparison":    "Expert Care level comparison: ",
        "sup_comparison_body": "Basic = 9×5 no fix-time  |  Advanced = 24×7 with fix-time SLA  |  Premium = 24×7 + dedicated Technical Account Manager",
        "cfg_model":         "Model",
        "cfg_form":          "Form Factor",
        "cfg_form_val":      "{ff} rack-mountable (19″ standard)",
        "cfg_data_nodes":    "Storage (Data) Nodes",
        "cfg_util_nodes":    "Utility Nodes",
        "cfg_drives":        "NVMe Drives (per data node)",
        "cfg_drive_type":    "Drive Type",
        "cfg_network":       "Network Interface",
        "cfg_fs":            "Parallel File System",
        "cfg_edition":       "Software Edition",
        "cfg_protocols":     "Supported Protocols",
        "cfg_enc":           "Data Encryption",
        "cfg_enc_yes":       "Enabled (Crypto-enabled network adapters)",
        "cfg_enc_no":        "Not configured",
        "cap_raw":           "Raw Capacity",
        "cap_usable":        "Usable Capacity",
        "cap_effective":     "Effective Capacity (with data reduction)",
        "cap_protection":    "Data Protection",
        "cap_drives_lbl":    "Drive Count",
        "perf_seq_read":     "Sequential Read Throughput",
        "perf_seq_write":    "Sequential Write Throughput",
        "perf_iops":         "Random IOPS",
        "perf_lat":          "Average Latency",
        "perf_gbs":          "GB/s",
        "perf_gibs":         "GiB/s",
        "perf_note":         "Note",
        "perf_note_val":     "Values computed by IBM Storage Performance Modeller. No guarantee implied.",
        "env_rack":          "Rack Space (per node)",
        "env_rack_val":      "{n}U (standard 19″ rack)",
        "env_pwr_typ":       "Power — Typical",
        "env_pwr_max":       "Power — Maximum",
        "env_pwr_val":       "{kw:.3f} kW  /  {kva:.3f} kVA",
        "env_cool":          "Cooling Requirement",
        "env_cool_val":      "{btu:,.0f} BTU/h",
        "env_psu":           "Power Supply",
        "env_psu_val":       "Dual redundant hot-swap PSUs (200–240V, 50/60 Hz)",
        "price_info":        "List prices from IBM e-config · Price file: {pf} · Discount applied: {d:.1f}% · Offer valid until: {vu}",
        "price_cat":         "Category",
        "price_qty":         "Qty",
        "price_list":        "List Price ({curr})",
        "price_disc":        "Discount",
        "price_eu_col":      "End User Price ({curr})",
        "price_hw":          "Hardware ({mc})",
        "price_sup":         "Expert Care Support",
        "price_sw":          "Software / Subscriptions",
        "price_ship":        "Shipping & Handling (non-discountable)",
        "price_total":       "TOTAL LIST PRICE",
        "price_eu_row":      "END USER PRICE",
        "price_fn":          (
            "* Prices are exclusive of applicable taxes. "
            "This offer is valid for 30 days from the date of preparation. "
            "Final pricing subject to IBM approval. "
            "Shipping and handling charges are non-discountable."
        ),
        "next_1_title":      "Technical Deep-Dive / POC",
        "next_1_body":       "Schedule a technical session with {client} to validate the proposed architecture against your workloads. IBM can provide a Proof of Concept environment on request.",
        "next_2_title":      "Formal Quotation",
        "next_2_body":       "Upon request, IBM or an authorised IBM Business Partner will issue a formal, binding quotation valid for 30 days, referencing the configuration ID from this document.",
        "next_3_title":      "Order Placement",
        "next_3_body":       "Orders are placed through an IBM Authorised Distributor or directly with IBM. Standard lead time for Storage Scale hardware is 6–8 weeks ARO (After Receipt of Order).",
        "next_4_title":      "Implementation & Onboarding",
        "next_4_body":       "IBM Lab Services or a certified IBM Business Partner will manage installation, network integration, file system configuration, and initial performance tuning.",
        "docs_heading":      "Product Documentation",
        "docs_ibm_docs":     "IBM Documentation",
        "docs_sales_manual": "IBM Sales Manual",
        "contact":           "Contact: {name} · IBM Storage Sales",
        "disclaimer":        (
            "This document is prepared for IBM Business Partners and their customers. "
            "Prices shown are list prices from IBM e-config and do not constitute a binding offer. "
            "Capacity values are calculated by IBM Storage Modeller — actual values may vary. "
            "Performance data is modelled and no guarantees are expressed or implied. "
            "IBM, Storage Scale, and GPFS are trademarks of International Business Machines Corporation."
        ),
    },
    "pl": {
        "cover_subtitle":    "Techniczne Podsumowanie Wykonawcze",
        "prepared_for":      "Przygotowane dla",
        "prepared_by":       "Przygotowane przez",
        "date":              "Data",
        "valid_until":       "Ważne do",
        "config_id":         "ID konfiguracji",
        "sec_exec":          "Podsumowanie Wykonawcze",
        "sec_config":        "Konfiguracja rozwiązania",
        "sec_capacity":      "Pojemność",
        "sec_performance":   "Profil wydajności",
        "sec_connectivity":  "Połączenia i ochrona danych",
        "sec_environment":   "Środowisko fizyczne",
        "sec_software":      "Oprogramowanie i subskrypcje",
        "sec_support":       "Serwis i wsparcie",
        "sec_pricing":       "Zestawienie cenowe",
        "sec_next":          "Kolejne kroki",
        "key_highlights":    "Kluczowe cechy rozwiązania",
        "platform_advantages": "Zalety platformy",
        "adv1": (
            "Równoległy system plików z globalną przestrzenią nazw — wszystkie węzły obliczeniowe "
            "uzyskują jednoczesny dostęp do tej samej przestrzeni, eliminując wąskie gardła I/O "
            "w potokach AI, HPC i analityki."
        ),
        "adv2": (
            "Ochrona danych z kodowaniem korekcyjnym (erasure coding) — parzystość rozłożona "
            "na wszystkich napędach klastra bez dedykowanych dysków spare, "
            "maksymalizując efektywność pojemności użytecznej."
        ),
        "adv3": (
            "Architektura active-active scale-out — dodawanie węzłów bez przerwy w pracy; "
            "wydajność i pojemność skalują się liniowo, a klienci kontynuują operacje I/O."
        ),
        "adv4": (
            "Przezroczyste tiering do chmury — nieaktywne dane migrują automatycznie "
            "do object storage (on-premises lub chmura), redukując zajętość pamięci podstawowej "
            "bez zmiany ścieżki dostępu dla aplikacji."
        ),
        "body": (
            "Niniejsza propozycja {client_str}rekomenduje rozwiązanie równoległego systemu plików {model_name}. "
            "Konfiguracja obejmuje {num_nodes} węzeł(ły) pamięci masowej z {drives_count} napędami NVMe "
            "({drive_type}), zapewniając {usable_tib:.1f} TiB pojemności użytecznej. "
            "Całość sprzętu objęta jest pakietem {support_name} w trybie {support_hours} przez {support_years} lat."
        ),
        "multi_system_note": "Niniejszy dokument obejmuje konfigurację {n} × {model}. Parametry poniżej dotyczą pojedynczego systemu.",
        "support_24x7":      "24×7 wsparcie całodobowe z fix-time SLA",
        "support_9x5":       "9×5 wsparcie w godzinach roboczych",
        "no_perf_data":      "Brak danych wydajnościowych. Wgraj raport wydajności ze Storage Modeller lub wpisz przepustowość ręcznie.",
        "no_support":        "Brak informacji o wsparciu w konfiguracji.",
        "sup_pkg":           "Pakiet wsparcia",
        "sup_level":         "Poziom",
        "sup_term":          "Okres",
        "sup_years_unit":    "{n} lat",
        "sup_coverage":      "Godziny dostępności",
        "sup_fixtime":       "Fix-time SLA",
        "sup_fixtime_yes":   "Tak — następny dzień roboczy lub ten sam dzień (na życzenie)",
        "sup_fixtime_no":    "Brak fix-time SLA",
        "sup_desc":          "Opis",
        "sup_comparison":    "Poziomy Expert Care: ",
        "sup_comparison_body": "Basic = 9×5 bez fix-time  |  Advanced = 24×7 z fix-time SLA  |  Premium = 24×7 + dedykowany Technical Account Manager",
        "cfg_model":         "Model",
        "cfg_form":          "Obudowa",
        "cfg_form_val":      "{ff} montaż w szafie rack (standard 19″)",
        "cfg_data_nodes":    "Węzły pamięci (Data Nodes)",
        "cfg_util_nodes":    "Węzły pomocnicze (Utility Nodes)",
        "cfg_drives":        "Napędy NVMe (na węzeł danych)",
        "cfg_drive_type":    "Typ napędu",
        "cfg_network":       "Interfejs sieciowy",
        "cfg_fs":            "Równoległy system plików",
        "cfg_edition":       "Edycja oprogramowania",
        "cfg_protocols":     "Obsługiwane protokoły",
        "cfg_enc":           "Szyfrowanie danych",
        "cfg_enc_yes":       "Włączone (adaptery sieciowe z obsługą kryptografii)",
        "cfg_enc_no":        "Nie skonfigurowano",
        "cap_raw":           "Pojemność fizyczna (raw)",
        "cap_usable":        "Pojemność użyteczna",
        "cap_effective":     "Pojemność efektywna (z redukcją danych)",
        "cap_protection":    "Ochrona danych",
        "cap_drives_lbl":    "Liczba napędów",
        "perf_seq_read":     "Przepustowość sekwencyjnego odczytu",
        "perf_seq_write":    "Przepustowość sekwencyjnego zapisu",
        "perf_iops":         "Losowe IOPS",
        "perf_lat":          "Średni czas odpowiedzi",
        "perf_gbs":          "GB/s",
        "perf_gibs":         "GiB/s",
        "perf_note":         "Uwaga",
        "perf_note_val":     "Wartości obliczone przez IBM Storage Performance Modeller. Bez gwarancji wyników.",
        "env_rack":          "Zajętość szafy (na węzeł)",
        "env_rack_val":      "{n}U (standardowa szafa 19″)",
        "env_pwr_typ":       "Pobór mocy — typowy",
        "env_pwr_max":       "Pobór mocy — maksymalny",
        "env_pwr_val":       "{kw:.3f} kW  /  {kva:.3f} kVA",
        "env_cool":          "Wymagania chłodzenia",
        "env_cool_val":      "{btu:,.0f} BTU/h",
        "env_psu":           "Zasilanie",
        "env_psu_val":       "Dwa nadmiarowe zasilacze hot-swap (200–240V, 50/60 Hz)",
        "price_info":        "Ceny katalogowe z IBM e-config · Plik cen: {pf} · Rabat: {d:.1f}% · Oferta ważna do: {vu}",
        "price_cat":         "Kategoria",
        "price_qty":         "Ilość",
        "price_list":        "Cena katalogowa ({curr})",
        "price_disc":        "Rabat",
        "price_eu_col":      "Cena dla klienta ({curr})",
        "price_hw":          "Sprzęt ({mc})",
        "price_sup":         "Expert Care Support",
        "price_sw":          "Oprogramowanie / Subskrypcje",
        "price_ship":        "Dostawa i obsługa (nie podlega rabatowi)",
        "price_total":       "ŁĄCZNA CENA KATALOGOWA",
        "price_eu_row":      "CENA DLA KLIENTA",
        "price_fn":          (
            "* Ceny nie zawierają podatków. "
            "Niniejsza oferta jest ważna przez 30 dni od daty sporządzenia. "
            "Ostateczna cena podlega zatwierdzeniu przez IBM. "
            "Opłaty za dostawę i obsługę nie podlegają rabacie."
        ),
        "next_1_title":      "Sesja techniczna / PoC",
        "next_1_body":       "Zaplanuj sesję techniczną z {client}, aby zweryfikować proponowaną architekturę względem rzeczywistych obciążeń. IBM może udostępnić środowisko Proof of Concept na żądanie.",
        "next_2_title":      "Formalna wycena",
        "next_2_body":       "Na żądanie IBM lub autoryzowany IBM Business Partner wystawi formalną, wiążącą wycenę ważną 30 dni, z odniesieniem do ID konfiguracji z niniejszego dokumentu.",
        "next_3_title":      "Złożenie zamówienia",
        "next_3_body":       "Zamówienia składane są przez autoryzowanego dystrybutora IBM lub bezpośrednio w IBM. Standardowy czas realizacji sprzętu Storage Scale wynosi 6–8 tygodni od przyjęcia zamówienia.",
        "next_4_title":      "Wdrożenie i uruchomienie",
        "next_4_body":       "IBM Lab Services lub certyfikowany IBM Business Partner zarządza instalacją, integracją sieciową, konfiguracją systemu plików i wstępnym tuningiem wydajnościowym.",
        "docs_heading":      "Dokumentacja produktu",
        "docs_ibm_docs":     "IBM Documentation",
        "docs_sales_manual": "IBM Sales Manual",
        "contact":           "Kontakt: {name} · IBM Storage Sales",
        "disclaimer":        (
            "Niniejszy dokument przygotowany jest dla IBM Business Partners i ich klientów. "
            "Prezentowane ceny są cenami katalogowymi z IBM e-config i nie stanowią wiążącej oferty. "
            "Wartości pojemności obliczane są przez IBM Storage Modeller — rzeczywiste wartości mogą się różnić. "
            "Dane wydajnościowe mają charakter modelowy i nie stanowią gwarancji. "
            "IBM, Storage Scale i GPFS są znakami towarowymi International Business Machines Corporation."
        ),
    },
}


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def generate_scale_exec_summary(
    project: dict[str, Any],
    client_name: str = "",
    seller_name: str = "",
    discount_pct: float = 60.0,
    lang: str = "en",
    num_systems: int = 1,
    eu_margin_pct: float = 15.0,
) -> bytes:
    """Generate Storage Scale Executive Summary DOCX and return as bytes."""
    T = _TRANS.get(lang, _TRANS["en"])

    doc = Document()
    _set_page_margins(doc)
    _set_default_font(doc)

    model_info = get_model_info(project.get("model_code", ""))
    pricing    = _calc_pricing(project, discount_pct, num_systems=num_systems, eu_margin_pct=eu_margin_pct)

    _add_cover_page(doc, project, model_info, client_name, seller_name, T)
    doc.add_page_break()

    _add_logo_header(doc)
    _add_exec_summary_text(doc, project, model_info, client_name, T, num_systems=num_systems)
    _add_section_heading(doc, T["sec_config"])
    _add_config_table(doc, project, model_info, T)
    _add_section_heading(doc, T["sec_capacity"])
    _add_capacity_table(doc, project, T)
    _add_section_heading(doc, T["sec_performance"])
    _add_performance_section(doc, project, T)
    _add_section_heading(doc, T["sec_environment"])
    _add_environment_table(doc, project, T)
    _add_section_heading(doc, T["sec_support"])
    _add_support_section(doc, project, T)
    _add_section_heading(doc, T["sec_pricing"])
    _add_pricing_table(doc, pricing, project, T)
    _add_section_heading(doc, T["sec_next"])
    _add_next_steps(doc, project, model_info, client_name, seller_name, T)
    _add_footer_disclaimer(doc, T)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _add_cover_page(doc, project, model_info, client_name, seller_name, T):
    from app.generators.exec_summary import (
        _add_logo_header, _set_cell_bg, _set_cell_text, _para_space,
        _get_logo_png, _convert_to_png,
    )

    _add_logo_header(doc)

    # Title block
    title_p = doc.add_paragraph()
    _para_space(title_p, before=40, after=4)
    run = title_p.add_run(T["cover_subtitle"])
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = IBM_DARK

    model_name = model_info.get("name", "IBM Storage Scale System")
    sub_p = doc.add_paragraph()
    _para_space(sub_p, before=0, after=24)
    run = sub_p.add_run(model_name)
    run.font.size = Pt(16)
    run.font.color.rgb = IBM_BLUE

    # Meta table
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    _date_fmt = "%d.%m.%Y" if T.get("date") == "Data" else "%B %d, %Y"
    today = date.today()
    valid = today + timedelta(days=30)

    meta_rows = [
        (T["prepared_for"], client_name or "—"),
        (T["prepared_by"],  seller_name or "—"),
        (T["date"],         today.strftime(_date_fmt)),
        (T["valid_until"],  valid.strftime(_date_fmt)),
    ]
    for i, (label, value) in enumerate(meta_rows):
        row = table.rows[i]
        _set_cell_bg(row.cells[0], IBM_DARK)
        _set_cell_text(row.cells[0], label, bold=True, color=IBM_WHITE, size=9)
        _set_cell_text(row.cells[1], value, size=10)


# ---------------------------------------------------------------------------
# Logo header
# ---------------------------------------------------------------------------

def _add_logo_header(doc):
    from app.generators.exec_summary import _add_logo_header as _fs_logo
    _fs_logo(doc)


# ---------------------------------------------------------------------------
# Executive Summary narrative
# ---------------------------------------------------------------------------

def _add_exec_summary_text(doc, project, model_info, client_name, T, num_systems: int = 1):
    from app.generators.exec_summary import _add_section_heading, _para_space

    _add_section_heading(doc, T["sec_exec"])

    model_name   = model_info.get("name", "IBM Storage Scale System")
    num_nodes    = project.get("num_data_nodes", 0) or project.get("num_nodes", 0) or 1
    drives_count = project.get("drives_count", 0)
    drive_type   = project.get("drive_type", "NVMe SSD")
    usable_tib   = project.get("usable_tib", 0.0)
    support_info = project.get("support_info") or {}
    support_name = support_info.get("name", "IBM Expert Care")
    support_years = support_info.get("years", 5)
    client_str   = f"for {client_name} " if client_name else ""

    if num_systems > 1:
        note_p = doc.add_paragraph()
        _para_space(note_p, before=0, after=8)
        run = note_p.add_run(T["multi_system_note"].format(n=num_systems, model=model_name))
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = IBM_BLUE

    support_hours = T["support_24x7"] if support_info.get("fix_time") else T["support_9x5"]
    body = T["body"].format(
        client_str=client_str,
        model_name=model_name,
        num_nodes=num_nodes,
        drives_count=drives_count,
        drive_type=drive_type,
        usable_tib=usable_tib,
        support_name=support_name,
        support_years=support_years,
        support_hours=support_hours,
    )
    p = doc.add_paragraph(body)
    _para_space(p, before=0, after=12)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = IBM_DARK

    # Highlights
    highlights = model_info.get("highlights", [])
    if highlights:
        hl_p = doc.add_paragraph()
        run = hl_p.add_run(T["key_highlights"])
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = IBM_BLUE
        _para_space(hl_p, before=6, after=2)
        for hl in highlights:
            bp = doc.add_paragraph(style="List Bullet")
            run = bp.add_run(hl)
            run.font.size = Pt(10)
            run.font.color.rgb = IBM_DARK
            _para_space(bp, before=0, after=1)

    # Platform advantages
    adv_p = doc.add_paragraph()
    run = adv_p.add_run(T["platform_advantages"])
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = IBM_BLUE
    _para_space(adv_p, before=8, after=2)

    for key in ("adv1", "adv2", "adv3", "adv4"):
        ap = doc.add_paragraph(style="List Bullet")
        run = ap.add_run(T[key])
        run.font.size = Pt(10)
        run.font.color.rgb = IBM_DARK
        _para_space(ap, before=0, after=2)


# ---------------------------------------------------------------------------
# Configuration table
# ---------------------------------------------------------------------------

def _add_config_table(doc, project, model_info, T):
    from app.generators.exec_summary import _make_two_col_table

    enc           = project.get("encryption", False)
    protocols     = project.get("protocol_support", [])
    proto_str     = ", ".join(protocols) if protocols else "NFS, SMB, S3"
    num_data      = project.get("num_data_nodes", 0) or project.get("num_nodes", 1)
    drives_node   = project.get("drives_per_node", 0) or project.get("drives_count", 0) // max(num_data, 1)
    edition       = project.get("scale_edition", "—")
    utility_nodes = project.get("utility_nodes", [])

    # Utility nodes summary string: "1 × Protocol Node (5149-23E), 1 × Management Server (5149-23E)"
    if utility_nodes:
        util_parts = []
        for u in utility_nodes:
            util_parts.append(f"{u['qty']} × {u['type']} ({u['mtm']})")
        util_str = ", ".join(util_parts)
    else:
        util_str = "—"

    rows = [
        (T["cfg_model"],      model_info.get("name", project.get("model_code", "—"))),
        (T["cfg_form"],       T["cfg_form_val"].format(ff=model_info.get("form_factor", "2U"))),
        (T["cfg_data_nodes"], str(num_data)),
        (T["cfg_util_nodes"], util_str),
        (T["cfg_drives"],     str(drives_node)),
        (T["cfg_drive_type"], project.get("drive_type", "—")),
        (T["cfg_network"],    project.get("network_type", "—")),
        (T["cfg_fs"],         project.get("filesystem_type", "IBM Storage Scale (GPFS)")),
        (T["cfg_edition"],    edition),
        (T["cfg_protocols"],  proto_str),
        (T["cfg_enc"],        T["cfg_enc_yes"] if enc else T["cfg_enc_no"]),
    ]
    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Capacity table
# ---------------------------------------------------------------------------

def _add_capacity_table(doc, project, T):
    from app.generators.exec_summary import _make_two_col_table

    raw_tib      = project.get("raw_tib", 0.0)
    raw_tb       = project.get("raw_tb", 0.0)
    usable_tib   = project.get("usable_tib", 0.0)
    usable_tb    = project.get("usable_tb", 0.0)
    effective_tib = project.get("effective_tib", 0.0)
    raid_type    = project.get("raid_type", "Erasure Code (8+2p)")
    rebuild_areas = project.get("rebuild_areas", 2)

    rows = [
        (T["cap_raw"],        f"{raw_tb:.2f} TB  /  {raw_tib:.2f} TiB"),
        (T["cap_usable"],     f"{usable_tb:.2f} TB  /  {usable_tib:.2f} TiB"),
        (T["cap_effective"],  f"{effective_tib:.2f} TiB"),
        (T["cap_protection"], f"{raid_type} — {rebuild_areas} fault domain(s)"),
        (T["cap_drives_lbl"], str(project.get("drives_count", "—"))),
    ]
    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Performance section
# ---------------------------------------------------------------------------

def _add_performance_section(doc, project, T):
    from app.generators.exec_summary import _make_two_col_table, _para_space

    tp_read_gbs  = project.get("throughput_read_gbs",  0.0)
    tp_write_gbs = project.get("throughput_write_gbs", 0.0)
    tp_read_gibs = project.get("throughput_read_gibs",  0.0)
    tp_write_gibs = project.get("throughput_write_gibs", 0.0)

    if not tp_read_gbs and not tp_write_gbs:
        p = doc.add_paragraph(T["no_perf_data"])
        _para_space(p, before=0, after=6)
        for run in p.runs:
            run.font.size = Pt(9)
        return

    rows = []
    if tp_read_gbs:
        rows.append((T["perf_seq_read"],
                     f"{tp_read_gbs:.2f} {T['perf_gbs']}  /  {tp_read_gibs:.2f} {T['perf_gibs']}"))
    if tp_write_gbs:
        rows.append((T["perf_seq_write"],
                     f"{tp_write_gbs:.2f} {T['perf_gbs']}  /  {tp_write_gibs:.2f} {T['perf_gibs']}"))
    rows.append((T["perf_note"], T["perf_note_val"]))

    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Environment table
# ---------------------------------------------------------------------------

def _add_environment_table(doc, project, T):
    from app.generators.exec_summary import _make_two_col_table

    rows = [
        (T["env_rack"],    T["env_rack_val"].format(n=project.get("rack_units", 2))),
        (T["env_pwr_typ"], T["env_pwr_val"].format(
            kw=project.get("power_kw_typical", 0.0),
            kva=project.get("power_kva_typical", 0.0))),
        (T["env_pwr_max"], T["env_pwr_val"].format(
            kw=project.get("power_kw_max", 0.0),
            kva=project.get("power_kva_max", 0.0))),
        (T["env_cool"],    T["env_cool_val"].format(btu=project.get("cooling_btu", 0.0))),
        (T["env_psu"],     T["env_psu_val"]),
    ]
    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Support section
# ---------------------------------------------------------------------------

def _add_support_section(doc, project, T):
    from app.generators.exec_summary import _make_two_col_table, _para_space

    support_info = project.get("support_info") or {}
    if not support_info:
        p = doc.add_paragraph(T["no_support"])
        return

    rows = [
        (T["sup_pkg"],      support_info.get("name", "—")),
        (T["sup_level"],    support_info.get("level", "—")),
        (T["sup_term"],     T["sup_years_unit"].format(n=support_info.get("years", "—"))),
        (T["sup_coverage"], support_info.get("coverage", "—")),
        (T["sup_fixtime"],  T["sup_fixtime_yes"] if support_info.get("fix_time") else T["sup_fixtime_no"]),
        (T["sup_desc"],     support_info.get("description", "")),
    ]
    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Pricing table (reuse FlashSystem pricing logic)
# ---------------------------------------------------------------------------

def _add_pricing_table(doc, pricing, project, T):
    from app.generators.exec_summary import _add_pricing_table as _fs_pricing
    _fs_pricing(doc, pricing, project, T)


# ---------------------------------------------------------------------------
# Next steps + documentation links
# ---------------------------------------------------------------------------

def _add_next_steps(doc, project, model_info, client_name, seller_name, T):
    from app.generators.exec_summary import _para_space

    _client = client_name or "your organisation"
    steps = [
        (T["next_1_title"], T["next_1_body"].format(client=_client)),
        (T["next_2_title"], T["next_2_body"]),
        (T["next_3_title"], T["next_3_body"]),
        (T["next_4_title"], T["next_4_body"]),
    ]
    for i, (title, body) in enumerate(steps):
        row_p = doc.add_paragraph()
        _para_space(row_p, before=4 if i > 0 else 0, after=0)
        num_run = row_p.add_run(f"{i+1}. ")
        num_run.font.bold = True
        num_run.font.size = Pt(10)
        num_run.font.color.rgb = IBM_BLUE
        title_run = row_p.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = IBM_DARK

        body_p = doc.add_paragraph(body)
        _para_space(body_p, before=1, after=4)
        for run in body_p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = IBM_GRAY

    # Contact
    contact_p = doc.add_paragraph()
    _para_space(contact_p, before=8, after=0)
    cr = contact_p.add_run(T["contact"].format(name=seller_name or "your IBM Sales Representative"))
    cr.font.size = Pt(9)
    cr.font.bold = True
    cr.font.color.rgb = IBM_BLUE

    # Docs links
    _short    = model_info.get("short", "")
    _docs     = get_docs(_short) if _short else {}
    _docs_url = _docs.get("docs_url", "")
    _sm_url   = _docs.get("sales_manual_url", "")
    if _docs_url or _sm_url:
        doc.add_paragraph()
        heading_p = doc.add_paragraph()
        _para_space(heading_p, before=8, after=2)
        hr = heading_p.add_run(T["docs_heading"])
        hr.font.size = Pt(9)
        hr.font.bold = True
        hr.font.color.rgb = IBM_DARK

        def _link_para(label, url):
            p = doc.add_paragraph()
            _para_space(p, before=1, after=1)
            lbl = p.add_run(f"{label}: ")
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = IBM_GRAY
            r_id = p.part.relate_to(url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True)
            hl = OxmlElement("w:hyperlink")
            hl.set(qn("r:id"), r_id)
            wr = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            style = OxmlElement("w:rStyle")
            style.set(qn("w:val"), "Hyperlink")
            rpr.append(style)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "18")
            rpr.append(sz)
            wr.append(rpr)
            wt = OxmlElement("w:t")
            wt.text = url
            wr.append(wt)
            hl.append(wr)
            p._p.append(hl)

        if _docs_url:
            _link_para(T["docs_ibm_docs"], _docs_url)
        if _sm_url:
            _link_para(T["docs_sales_manual"], _sm_url)


# ---------------------------------------------------------------------------
# Footer disclaimer
# ---------------------------------------------------------------------------

def _add_footer_disclaimer(doc, T):
    from app.generators.exec_summary import _add_hrule, _para_space

    doc.add_paragraph()
    _add_hrule(doc)
    p = doc.add_paragraph()
    _para_space(p, before=4, after=0)
    run = p.add_run(T["disclaimer"])
    run.font.size = Pt(8)
    run.font.color.rgb = IBM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Pricing calculation (reuse FlashSystem helper)
# ---------------------------------------------------------------------------

def _calc_pricing(project, discount_pct, num_systems=1, eu_margin_pct=15.0):
    from app.generators.exec_summary import _calc_pricing as _fs_calc
    return _fs_calc(project, discount_pct, num_systems=num_systems, eu_margin_pct=eu_margin_pct)


# ---------------------------------------------------------------------------
# DOCX helpers (reuse FlashSystem helpers)
# ---------------------------------------------------------------------------

def _set_page_margins(doc):
    from app.generators.exec_summary import _set_page_margins as _fs_margins
    _fs_margins(doc)


def _set_default_font(doc):
    from app.generators.exec_summary import _set_default_font as _fs_font
    _fs_font(doc)


def _add_section_heading(doc, text):
    from app.generators.exec_summary import _add_section_heading as _fs_heading
    _fs_heading(doc, text)
