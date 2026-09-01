"""
Executive Summary DOCX generator.
Produces a professional IBM-branded Word document from a parsed project dict.
"""

from __future__ import annotations

import io
import math
import re
from datetime import date, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Inches, Emu

from app.knowledge.product_db import get_model_info, get_feature_label, get_docs

# IBM design tokens
IBM_BLUE = RGBColor(0x00, 0x62, 0xFF)
IBM_DARK = RGBColor(0x16, 0x16, 0x16)
IBM_GRAY = RGBColor(0x52, 0x52, 0x52)
IBM_LIGHT_GRAY = RGBColor(0xF4, 0xF4, 0xF4)
IBM_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ASSETS_DIR = Path(__file__).parent.parent / "assets"
LOGOS_DIR = ASSETS_DIR / "logos"
IMAGES_DIR = ASSETS_DIR / "images"

# SAS-based models (FS5045/FS5015) — no FlashCore Modules, no hardware-accelerated
# compression/dedup/ransomware-detection, 16 Gb FC (ALBG), SAS internal bus
_SAS_MODELS: frozenset[str] = frozenset({"FS5045", "FS5015"})

# ISO 3166-1 alpha-2 → English country name (European + common IBM markets)
_COUNTRY_NAMES: dict[str, str] = {
    "AT": "Austria",        "BE": "Belgium",        "BG": "Bulgaria",
    "CH": "Switzerland",    "CY": "Cyprus",          "CZ": "Czech Republic",
    "DE": "Germany",        "DK": "Denmark",         "EE": "Estonia",
    "ES": "Spain",          "FI": "Finland",         "FR": "France",
    "GB": "United Kingdom", "GR": "Greece",          "HR": "Croatia",
    "HU": "Hungary",        "IE": "Ireland",         "IS": "Iceland",
    "IT": "Italy",          "LI": "Liechtenstein",   "LT": "Lithuania",
    "LU": "Luxembourg",     "LV": "Latvia",          "MT": "Malta",
    "NL": "Netherlands",    "NO": "Norway",           "PL": "Poland",
    "PT": "Portugal",       "RO": "Romania",          "SE": "Sweden",
    "SI": "Slovenia",       "SK": "Slovakia",         "TR": "Turkey",
    "UA": "Ukraine",        "AE": "UAE",              "SA": "Saudi Arabia",
    "ZA": "South Africa",   "US": "United States",    "CA": "Canada",
    "AU": "Australia",      "NZ": "New Zealand",      "SG": "Singapore",
    "JP": "Japan",          "KR": "South Korea",      "IN": "India",
    "BR": "Brazil",         "MX": "Mexico",           "IL": "Israel",
}


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Translations
# ---------------------------------------------------------------------------

_TRANS: dict[str, dict[str, str]] = {
    "en": {
        "lang":                 "en",
        "cover_subtitle":       "Technical Executive Summary",
        "prepared_for":         "Prepared for",
        "prepared_by":          "Prepared by",
        "date":                 "Date",
        "valid_until":          "Valid until",
        "config_id":            "Configuration ID",
        "sec_exec":             "Executive Summary",
        "sec_config":           "Solution Configuration",
        "sec_capacity":         "Capacity Overview",
        "sec_performance":      "Performance Profile",
        "sec_connectivity":     "Connectivity & Data Protection",
        "sec_environment":      "Physical Environment",
        "sec_software":         "Software & Subscriptions",
        "sec_support":          "Service & Support",
        "sec_pricing":          "Pricing Summary",
        "sec_next":             "Next Steps",
        "key_highlights":       "Key Solution Highlights",
        "platform_advantages":  "IBM FlashSystem Platform Advantages",
        "adv1": (
            "Hardware-accelerated inline compression and deduplication "
            "— data reduction runs entirely at the drive level, eliminating CPU overhead and preserving full array throughput."
        ),
        "adv2": (
            "Hardware-accelerated ransomware detection — AI/ML anomaly detection executes directly "
            "in every NVMe drive with detection latency below 60 seconds, without any impact on host I/O performance."
        ),
        "adv3": (
            "Policy-based active-active HA replication — "
            "synchronous, zero-RPO replication with transparent failover, "
            "eliminating single points of failure across sites. "
            "No scale-out cluster required; HA is enforced through IBM Storage Virtualize replication policies."
        ),
        "adv4": (
            "WORM snapshots (immutable copies) — snapshot policies can enforce write-once/read-many "
            "retention to protect against accidental deletion, ransomware encryption, and insider "
            "threats, supporting regulatory compliance requirements."
        ),
        "body": (
            "The proposed solution {client_str}is based on the {model_name} — "
            "an enterprise all-NVMe storage system designed for demanding, performance-sensitive workloads. "
            "The configuration includes {drives_count} × {drive_type} drives delivering "
            "{usable_tib:.2f} TiB of usable capacity and up to {effective_tb:.2f} TB of effective capacity "
            "with inline data reduction enabled.\n\n"
            "The system provides hardware-accelerated AI-powered ransomware detection at the drive level "
            "— with zero performance penalty. "
            "All data at rest is protected by software encryption, and high availability is ensured by "
            "a Distributed RAID 6 configuration with {cache_gb} GB of cache memory.\n\n"
            "The solution is covered by {support_name} ({support_years}-year term), providing "
            "{support_hours}."
        ),
        "support_24x7":         "24×7 around-the-clock support with a hardware fix-time SLA",
        "support_9x5":          "9×5 business-hours support",
        "no_perf_data":         "Performance data not provided. Upload a Storage Modeller performance report or enter IOPS manually.",
        "no_support":           "No support information found in configuration.",
        "sup_pkg":              "Support Package",
        "sup_level":            "Level",
        "sup_term":             "Term",
        "sup_years_unit":       "{n} years",
        "sup_coverage":         "Coverage Hours",
        "sup_fixtime":          "Hardware Fix-Time SLA",
        "sup_fixtime_yes":      "Yes — {hours}",
        "sup_fixtime_no":       "No fix-time SLA",
        "sup_desc":             "Description",
        "sup_comparison":       "Expert Care level comparison: ",
        "sup_comparison_body":  "Basic = 9×5 no fix-time  |  Advanced = 24×7 with fix-time SLA  |  Premium = 24×7 + dedicated Technical Account Manager",
        "cfg_model":            "Model",
        "cfg_form":             "Form Factor",
        "cfg_form_val":         "{ff} rack-mountable (19\u2033 standard)",
        "cfg_fw":               "Firmware Version",
        "cfg_iogroups":         "I/O Groups",
        "cfg_enclosures":       "Enclosures",
        "cfg_drives":           "Flash Drives",
        "cfg_cache":            "Cache Memory",
        "cfg_cache_val":        "{n} GB per I/O group",
        "cfg_host":             "Host Connectivity",
        "cfg_host_fc":          "{n} × 32 Gb Fibre Channel ports",
        "cfg_host_eth":         "On-board 25/10 GbE",
        "cfg_enc_yes":          "Software encryption enabled (ACEG)",
        "cfg_enc_no":           "Not configured",
        "cfg_encryption":       "Encryption",
        "cfg_drive_slots":      "Drive Slots",
        "cfg_drive_slots_val":  "{total} total · {used} installed · {free} available (hot-add online)",
        "cap_raw":              "Raw Capacity",
        "cap_usable":           "Usable Capacity (after RAID)",
        "cap_effective":        "Effective Capacity (with data reduction)",
        "cap_rec":              "Recommended Max Usage",
        "cap_rec_val":          "{n:.0f} TiB (80% of usable)",
        "cap_raid":             "RAID Configuration",
        "cap_raid_val":         "{raid} — {areas} rebuild area(s), tolerates 2 simultaneous drive failures",
        "cap_drives":           "Drive Count",
        "cap_drives_val":       "{n} drives in 1 pool · 1 array",
        "perf_profile":         "Workload Profile",
        "perf_profile_val":     "{read:.0f}% Read / {write:.0f}% Write · {bs} KiB block size",
        "perf_total":           "Total I/O Rate",
        "perf_read":            "Read I/O Rate",
        "perf_write":           "Write I/O Rate",
        "perf_tp":              "Throughput",
        "perf_tp_val":          "{mib:,.1f} MiB/s  ({gib:.2f} GiB/s)",
        "perf_lat":             "Average Response Time",
        "perf_lat_val":         "{ms:.3f} ms/op",
        "perf_cache":           "Cache Read Hit Rate",
        "perf_note":            "Note",
        "perf_note_val":        "Values computed by IBM Storage Performance Modeller. No guarantee implied.",
        "conn_fc":              "Host Interface — Fibre Channel",
        "conn_fc_val":          "{n} × 32 Gb FC ports (ALB9 adapter pair)",
        "conn_fc_no":           "Not configured",
        "conn_eth":             "Host Interface — Ethernet",
        "conn_eth_val":         "4 × 25/10 GbE on-board ports (NVMe-oF / iSCSI)",
        "conn_cables":          "FC Cables Included",
        "conn_cables_val":      "{n} × 5 m OM3 LC fibre cables",
        "conn_cables_no":       "Not included",
        "conn_mgmt":            "Management Interface",
        "conn_mgmt_val":        "Dedicated 1 GbE management ports (separate from host I/O)",
        "conn_bus":             "Internal Bus",
        "conn_bus_val":         "PCIe 4.0 (NVMe drives to controllers)",
        "conn_drv_prot":        "Drive Protection",
        "conn_enc":             "Data Encryption",
        "conn_enc_yes":         "Software-based full encryption (ACEG) — no performance impact",
        "conn_enc_no":          "Not configured",
        "conn_ransom":          "Ransomware Detection",
        "conn_ransom_val":      "AI-powered hardware-accelerated per-drive ransomware detection (no performance penalty)",
        "env_rack":             "Rack Space",
        "env_rack_val":         "{n}U (standard 19\u2033 rack)",
        "env_pwr_typ":          "Power — Typical",
        "env_pwr_max":          "Power — Maximum",
        "env_pwr_val":          "{kw:.3f} kW  /  {kva:.3f} kVA",
        "env_cool":             "Cooling Requirement",
        "env_cool_val":         "{btu:,.0f} BTU/h",
        "env_psu":              "Power Supply",
        "env_psu_val":          "Dual redundant hot-swap PSUs (200–240V, 50/60 Hz)",
        "env_ff":               "Form Factor",
        "env_ff_val":           "1U NVMe Control Enclosure",
        "sw_code":              "Product Code",
        "sw_desc":              "Description",
        "multi_system_note":    "This document covers a configuration of {n} × {model}. Performance and capacity specifications below refer to a single system.",
        "price_info":           "List prices from IBM e-config · Price file: {pf} · Discount applied: {d:.1f}% · Offer valid until: {vu}",
        "price_cat":            "Category",
        "price_qty":            "Qty",
        "price_list":           "List Price ({curr})",
        "price_disc":           "Discount",
        "price_eu_col":         "End User Price ({curr})",
        "price_hw":             "Hardware ({mc} + options)",
        "price_sup":            "Expert Care Support (5132-A05)",
        "price_sw":             "Software OTC",
        "price_ship":           "Shipping & Handling (non-discountable)",
        "price_total":          "TOTAL LIST PRICE",
        "price_eu_row":         "END USER PRICE",
        "price_fn":             (
            "* Prices are exclusive of applicable taxes. "
            "This offer is valid for 30 days from the date of preparation. "
            "Final pricing subject to IBM approval. "
            "Shipping and handling charges are non-discountable."
        ),
        "next_1_title":         "Technical Deep-Dive / POC",
        "next_1_body":          "Schedule a technical session with {client} to validate the proposed architecture against your workloads. IBM can provide a Proof of Concept environment on request.",
        "next_2_title":         "Formal Quotation",
        "next_2_body":          "Upon request, IBM or an authorised IBM Business Partner will issue a formal, binding quotation valid for 30 days, referencing the configuration ID from this document.",
        "next_3_title":         "Order Placement",
        "next_3_body":          "Orders are placed through an IBM Authorised Distributor or directly with IBM. Standard lead time for FlashSystem hardware is 4–6 weeks ARO (After Receipt of Order).",
        "next_4_title":         "Implementation & Onboarding",
        "next_4_body":          "IBM Lab Services or a certified IBM Business Partner will manage installation, data migration planning, and initial configuration. IBM Storage Insights is activated automatically upon first connection.",
        "docs_heading":         "Product Documentation",
        "docs_ibm_docs":        "IBM Documentation",
        "docs_sales_manual":    "IBM Sales Manual",
        "contact":              "Contact: {name} · IBM Storage Sales",
        "disclaimer":           (
            "This document is prepared for IBM Business Partners and their customers. "
            "Prices shown are list prices from IBM e-config and do not constitute a binding offer. "
            "Capacity values are calculated by IBM Storage Modeller — actual values may vary. "
            "Performance data is modelled and no guarantees are expressed or implied. "
            "IBM, FlashSystem, and FlashCore are trademarks of International Business Machines Corporation."
        ),
    },
    "pl": {
        "lang":                 "pl",
        "cover_subtitle":       "Techniczne Podsumowanie Wykonawcze",
        "prepared_for":         "Przygotowane dla",
        "prepared_by":          "Przygotowane przez",
        "date":                 "Data",
        "valid_until":          "Ważne do",
        "config_id":            "ID konfiguracji",
        "sec_exec":             "Podsumowanie Wykonawcze",
        "sec_config":           "Konfiguracja rozwiązania",
        "sec_capacity":         "Pojemność",
        "sec_performance":      "Profil wydajności",
        "sec_connectivity":     "Połączenia i ochrona danych",
        "sec_environment":      "Środowisko fizyczne",
        "sec_software":         "Oprogramowanie i subskrypcje",
        "sec_support":          "Serwis i wsparcie",
        "sec_pricing":          "Zestawienie cenowe",
        "sec_next":             "Kolejne kroki",
        "key_highlights":       "Kluczowe cechy rozwiązania",
        "platform_advantages":  "Przewagi platformy IBM FlashSystem",
        "adv1": (
            "Sprzętowa kompresja i deduplikacja inline "
            "— redukcja danych realizowana całkowicie na poziomie napędu, bez obciążenia CPU i bez wpływu na wydajność macierzy."
        ),
        "adv2": (
            "Sprzętowe wykrywanie ransomware — wnioskowanie AI/ML realizowane bezpośrednio "
            "w każdym napędzie NVMe; czas wykrycia anomalii poniżej 60 sekund, bez wpływu na operacje I/O hosta."
        ),
        "adv3": (
            "Replikacja active-active HA oparta na politykach — "
            "replikacja synchroniczna RPO=0 z transparentnym przełączeniem awaryjnym, "
            "eliminująca pojedyncze punkty awarii między lokalizacjami. "
            "Nie wymaga klastra scale-out; HA realizowana przez polityki replikacji IBM Storage Virtualize."
        ),
        "adv4": (
            "Migawki WORM (kopie niezmienne) — polityki migawek wymuszają retencję zapisu "
            "write-once/read-many, chroniąc przed przypadkowym usunięciem, szyfrowaniem przez "
            "ransomware i zagrożeniami wewnętrznymi, wspierając wymogi compliance."
        ),
        "body": (
            "Proponowane rozwiązanie {client_str}opiera się na macierzy {model_name} — "
            "korporacyjnym systemie pamięci masowej all-NVMe zaprojektowanym do wymagających "
            "obciążeń wrażliwych na wydajność. "
            "Konfiguracja zawiera {drives_count} × {drive_type}, zapewniając "
            "{usable_tib:.2f} TiB pojemności użytecznej i do {effective_tb:.2f} TB pojemności efektywnej "
            "przy włączonej redukcji danych inline.\n\n"
            "System oferuje sprzętowe, przyspieszone przez AI wykrywanie ransomware na poziomie napędu "
            "— bez negatywnego wpływu na wydajność. "
            "Wszystkie dane w spoczynku są chronione przez programowe szyfrowanie, "
            "a wysoka dostępność jest zapewniona przez konfigurację Distributed RAID 6 "
            "z {cache_gb} GB pamięci podręcznej.\n\n"
            "Rozwiązanie objęte jest pakietem {support_name} (umowa {support_years}-letnia), "
            "zapewniającym {support_hours}."
        ),
        "support_24x7":         "wsparcie całodobowe 24×7 z SLA fix-time dla sprzętu",
        "support_9x5":          "wsparcie w godzinach roboczych 9×5",
        "no_perf_data":         "Brak danych wydajnościowych. Prześlij raport Storage Modeller lub wprowadź IOPS ręcznie.",
        "no_support":           "Brak informacji o pakiecie serwisowym w konfiguracji.",
        "sup_pkg":              "Pakiet serwisowy",
        "sup_level":            "Poziom",
        "sup_term":             "Okres umowy",
        "sup_years_unit":       "{n} lat",
        "sup_coverage":         "Godziny wsparcia",
        "sup_fixtime":          "SLA fix-time dla sprzętu",
        "sup_fixtime_yes":      "Tak — {hours}",
        "sup_fixtime_no":       "Brak SLA fix-time",
        "sup_desc":             "Opis",
        "sup_comparison":       "Porównanie poziomów Expert Care: ",
        "sup_comparison_body":  "Basic = 9×5 bez fix-time  |  Advanced = 24×7 z SLA fix-time  |  Premium = 24×7 + dedykowany Technical Account Manager",
        "cfg_model":            "Model",
        "cfg_form":             "Współczynnik obudowy",
        "cfg_form_val":         "{ff} montaż w szafie 19\u2033",
        "cfg_fw":               "Wersja firmware",
        "cfg_iogroups":         "Grupy I/O",
        "cfg_enclosures":       "Obudowy",
        "cfg_drives":           "Napędy Flash",
        "cfg_cache":            "Pamięć cache",
        "cfg_cache_val":        "{n} GB na grupę I/O",
        "cfg_host":             "Połączenie z hostami",
        "cfg_host_fc":          "{n} × 32 Gb Fibre Channel",
        "cfg_host_eth":         "Porty 25/10 GbE wbudowane",
        "cfg_enc_yes":          "Włączone szyfrowanie programowe (ACEG)",
        "cfg_enc_no":           "Nie skonfigurowano",
        "cfg_encryption":       "Szyfrowanie",
        "cfg_drive_slots":      "Sloty dyskowe",
        "cfg_drive_slots_val":  "{total} łącznie · {used} zainstalowanych · {free} wolnych (dodawalne online)",
        "cap_raw":              "Pojemność fizyczna (raw)",
        "cap_usable":           "Pojemność użyteczna (po RAID)",
        "cap_effective":        "Pojemność efektywna (z redukcją danych)",
        "cap_rec":              "Zalecane maks. wykorzystanie",
        "cap_rec_val":          "{n:.0f} TiB (80% pojemności użytecznej)",
        "cap_raid":             "Konfiguracja RAID",
        "cap_raid_val":         "{raid} — {areas} obszar(y) odbudowy, toleruje 2 jednoczesne awarie napędów",
        "cap_drives":           "Liczba napędów",
        "cap_drives_val":       "{n} napędów w 1 puli · 1 macierz",
        "perf_profile":         "Profil obciążenia",
        "perf_profile_val":     "{read:.0f}% odczyt / {write:.0f}% zapis · rozmiar bloku {bs} KiB",
        "perf_total":           "Łączna wydajność I/O",
        "perf_read":            "Wydajność odczytu",
        "perf_write":           "Wydajność zapisu",
        "perf_tp":              "Przepustowość",
        "perf_tp_val":          "{mib:,.1f} MiB/s  ({gib:.2f} GiB/s)",
        "perf_lat":             "Średni czas odpowiedzi",
        "perf_lat_val":         "{ms:.3f} ms/op",
        "perf_cache":           "Współczynnik trafień cache (odczyt)",
        "perf_note":            "Uwaga",
        "perf_note_val":        "Wartości obliczone przez IBM Storage Performance Modeller. Bez gwarancji wyników.",
        "conn_fc":              "Interfejs hosta — Fibre Channel",
        "conn_fc_val":          "{n} × 32 Gb FC (para adapterów ALB9)",
        "conn_fc_no":           "Nie skonfigurowano",
        "conn_eth":             "Interfejs hosta — Ethernet",
        "conn_eth_val":         "4 × 25/10 GbE porty wbudowane (NVMe-oF / iSCSI)",
        "conn_cables":          "Kable FC w zestawie",
        "conn_cables_val":      "{n} × 5 m OM3 LC kable światłowodowe",
        "conn_cables_no":       "Nie w zestawie",
        "conn_mgmt":            "Interfejs zarządzania",
        "conn_mgmt_val":        "Dedykowane porty 1 GbE zarządzania (oddzielone od I/O hosta)",
        "conn_bus":             "Magistrala wewnętrzna",
        "conn_bus_val":         "PCIe 4.0 (napędy NVMe do kontrolerów)",
        "conn_drv_prot":        "Ochrona napędów",
        "conn_enc":             "Szyfrowanie danych",
        "conn_enc_yes":         "Pełne szyfrowanie programowe (ACEG) — bez wpływu na wydajność",
        "conn_enc_no":          "Nie skonfigurowano",
        "conn_ransom":          "Wykrywanie ransomware",
        "conn_ransom_val":      "Sprzętowe, przyspieszone sprzętowo wykrywanie ransomware per-napęd (bez kary wydajnościowej)",
        "env_rack":             "Zajętość szafy",
        "env_rack_val":         "{n}U (standardowa szafa 19\u2033)",
        "env_pwr_typ":          "Pobór mocy — typowy",
        "env_pwr_max":          "Pobór mocy — maksymalny",
        "env_pwr_val":          "{kw:.3f} kW  /  {kva:.3f} kVA",
        "env_cool":             "Wymagania chłodzenia",
        "env_cool_val":         "{btu:,.0f} BTU/h",
        "env_psu":              "Zasilanie",
        "env_psu_val":          "Dwa nadmiarowe zasilacze hot-swap (200–240V, 50/60 Hz)",
        "env_ff":               "Obudowa",
        "env_ff_val":           "1U obudowa sterowania NVMe",
        "sw_code":              "Kod produktu",
        "sw_desc":              "Opis",
        "multi_system_note":    "Niniejszy dokument obejmuje konfigurację {n} × {model}. Parametry wydajnościowe i pojemnościowe poniżej dotyczą pojedynczego systemu.",
        "price_info":           "Ceny katalogowe z IBM e-config · Plik cen: {pf} · Rabat: {d:.1f}% · Oferta ważna do: {vu}",
        "price_cat":            "Kategoria",
        "price_qty":            "Ilość",
        "price_list":           "Cena katalogowa ({curr})",
        "price_disc":           "Rabat",
        "price_eu_col":         "Cena dla klienta ({curr})",
        "price_hw":             "Sprzęt ({mc} + opcje)",
        "price_sup":            "Expert Care Support (5132-A05)",
        "price_sw":             "Oprogramowanie OTC",
        "price_ship":           "Dostawa i obsługa (nie podlega rabatowi)",
        "price_total":          "ŁĄCZNA CENA KATALOGOWA",
        "price_eu_row":         "CENA DLA KLIENTA",
        "price_fn":             (
            "* Ceny nie zawierają podatków. "
            "Niniejsza oferta jest ważna przez 30 dni od daty sporządzenia. "
            "Ostateczna cena podlega zatwierdzeniu przez IBM. "
            "Opłaty za dostawę i obsługę nie podlegają rabacie."
        ),
        "next_1_title":         "Sesja techniczna / PoC",
        "next_1_body":          "Zaplanuj sesję techniczną z {client}, aby zweryfikować proponowaną architekturę względem rzeczywistych obciążeń. IBM może udostępnić środowisko Proof of Concept na żądanie.",
        "next_2_title":         "Formalna wycena",
        "next_2_body":          "Na żądanie IBM lub autoryzowany IBM Business Partner wystawi formalną, wiążącą wycenę ważną 30 dni, z odniesieniem do ID konfiguracji z niniejszego dokumentu.",
        "next_3_title":         "Złożenie zamówienia",
        "next_3_body":          "Zamówienia składane są przez autoryzowanego dystrybutora IBM lub bezpośrednio w IBM. Standardowy czas realizacji sprzętu FlashSystem wynosi 4–6 tygodni od przyjęcia zamówienia.",
        "next_4_title":         "Wdrożenie i uruchomienie",
        "next_4_body":          "IBM Lab Services lub certyfikowany IBM Business Partner zarządza instalacją, planowaniem migracji danych i wstępną konfiguracją. IBM Storage Insights aktywuje się automatycznie po pierwszym podłączeniu.",
        "docs_heading":         "Dokumentacja produktu",
        "docs_ibm_docs":        "IBM Documentation",
        "docs_sales_manual":    "IBM Sales Manual",
        "contact":              "Kontakt: {name} · IBM Storage Sales",
        "disclaimer":           (
            "Niniejszy dokument przygotowany jest dla IBM Business Partners i ich klientów. "
            "Prezentowane ceny są cenami katalogowymi z IBM e-config i nie stanowią wiążącej oferty. "
            "Wartości pojemności obliczane są przez IBM Storage Modeller — rzeczywiste wartości mogą się różnić. "
            "Dane wydajnościowe mają charakter modelowy i nie stanowią gwarancji. "
            "IBM, FlashSystem i FlashCore są znakami towarowymi International Business Machines Corporation."
        ),
    },
}


def generate_exec_summary(
    project: dict[str, Any],
    client_name: str = "",
    seller_name: str = "",
    discount_pct: float = 60.0,
    iops_override: int | None = None,
    lang: str = "en",
    num_systems: int = 1,
    eu_margin_pct: float = 15.0,
) -> bytes:
    """
    Generate Exec Summary DOCX and return as bytes.

    Args:
        project: unified dict from parse_project().
        client_name: end-customer name.
        seller_name: sales rep name.
        discount_pct: discount percentage (0–100).
        iops_override: manual IOPS if no performance file was uploaded.
        num_systems: number of systems (multiplies price, adds note when > 1).
        eu_margin_pct: partner margin — BP = EU × (1 - margin%) (default 15%).

    Returns:
        DOCX file content as bytes.
    """
    T = _TRANS.get(lang, _TRANS["en"])

    doc = Document()
    _set_page_margins(doc)
    _set_default_font(doc)

    model_info = get_model_info(project.get("model_code", ""))
    pricing = _calc_pricing(project, discount_pct, num_systems=num_systems, eu_margin_pct=eu_margin_pct)
    iops = iops_override if iops_override else project.get("perf_iops_total", 0)

    # ── Page 1: Cover ────────────────────────────────────────────────────────
    _add_cover_page(doc, project, model_info, client_name, seller_name, T)

    # ── Page 2+: Content (no blank page between cover and content) ───────────
    _add_logo_header(doc)
    _add_executive_summary_text(doc, project, model_info, client_name, T, num_systems=num_systems)
    _add_section_heading(doc, T["sec_config"])
    _add_config_table(doc, project, model_info, T)
    _add_section_heading(doc, T["sec_capacity"])
    _add_capacity_table(doc, project, T)
    _add_section_heading(doc, T["sec_performance"])
    _add_performance_table(doc, project, iops, T)
    _add_section_heading(doc, T["sec_connectivity"])
    _add_connectivity_table(doc, project, model_info, T)
    _add_section_heading(doc, T["sec_environment"])
    _add_environment_table(doc, project, T)
    _add_section_heading(doc, T["sec_software"])
    _add_software_table(doc, project, T)
    _add_section_heading(doc, T["sec_support"])
    _add_support_section(doc, project, T)
    _add_section_heading(doc, T["sec_pricing"])
    _add_pricing_table(doc, pricing, project, T)
    _add_section_heading(doc, T["sec_next"])
    _add_next_steps(doc, project, model_info, client_name, seller_name, T)
    _add_footer_disclaimer(doc, T)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _add_cover_page(doc, project, model_info, client_name, seller_name, T):
    """
    Premium cover page with:
    - Clean white header with IBM 2026 logo (left) + blue accent line
    - Large centered product image
    - Two-line title: "Executive Summary" + model name
    - Blue accent metadata table
    """
    # ═══ 1. HEADER ROW: IBM LOGO (left) ═══
    new_logo = LOGOS_DIR / "ibm-logo-2026.png"
    if new_logo.exists():
        hdr_p = doc.add_paragraph()
        hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _para_space(hdr_p, before=0, after=6)
        hdr_run = hdr_p.add_run()
        hdr_run.add_picture(str(new_logo), width=Cm(4))

    # ═══ 2. BLUE ACCENT LINE ═══
    accent_p = doc.add_paragraph()
    _para_space(accent_p, before=0, after=0)
    pPr = accent_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top_b = OxmlElement("w:top")
    top_b.set(qn("w:val"), "single")
    top_b.set(qn("w:sz"), "18")  # 9pt thick
    top_b.set(qn("w:space"), "0")
    top_b.set(qn("w:color"), "0062FF")
    pBdr.append(top_b)
    pPr.append(pBdr)
    
    # ═══ 3. PRODUCT IMAGE (large, centered) ═══
    image_file = model_info.get("image")
    if image_file:
        direct_path = IMAGES_DIR / image_file
        if direct_path.exists() and direct_path.suffix.lower() == ".png":
            image_path = direct_path
        else:
            png_name = Path(image_file).stem + ".png"
            image_path = IMAGES_DIR / png_name
            if not image_path.exists():
                webp_path = IMAGES_DIR / image_file
                if webp_path.exists():
                    _convert_to_png(webp_path, image_path)
        
        if image_path.exists():
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _para_space(img_p, before=10, after=10)
            img_run = img_p.add_run()
            img_run.add_picture(str(image_path), width=Cm(10))
    
    # ═══ 4. TWO-LINE TITLE ═══
    # Line 1: "Executive Summary" (gray, 14pt)
    doc_type_p = doc.add_paragraph()
    doc_type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(doc_type_p, before=0, after=2)
    dt_run = doc_type_p.add_run(T.get("cover_title", "Executive Summary"))
    dt_run.font.size = Pt(14)
    dt_run.font.color.rgb = IBM_GRAY
    dt_run.font.name = "Calibri Light"
    
    # Line 2: Model name (blue, bold, 32pt)
    model_p = doc.add_paragraph()
    model_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(model_p, before=0, after=4)
    model_run = model_p.add_run(model_info.get("name", "IBM Storage Solution"))
    model_run.font.size = Pt(32)
    model_run.font.bold = True
    model_run.font.color.rgb = IBM_BLUE
    
    # Subtitle: model code (small gray)
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(sub_p, before=0, after=40)
    sub_run = sub_p.add_run(f"{T.get('cover_subtitle', 'Technical Overview')} · {project.get('model_code', '')}")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = IBM_GRAY
    
    # ═══ 5. METADATA TABLE WITH BLUE LEFT ACCENT ═══
    # Blue vertical accent bar (1-col table)
    accent_table = doc.add_table(rows=1, cols=2)
    accent_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(accent_table)
    
    accent_bar = accent_table.rows[0].cells[0]
    accent_bar.width = Cm(0.4)
    _set_cell_bg(accent_bar, "0062FF")
    
    # Metadata cell
    meta_cell = accent_table.rows[0].cells[1]
    meta_cell.width = Cm(14)
    
    today = date.today()
    valid_until = today + timedelta(days=30)
    _date_fmt = "%d.%m.%Y" if T.get("date") == "Data" else "%B %d, %Y"
    
    _country_code = project.get("country_code", "")
    _country_name = _COUNTRY_NAMES.get(_country_code, _country_code)
    _client_with_country = (
        f"{client_name}, {_country_name}" if client_name and _country_name
        else (client_name or "—")
    )
    
    rows_data = [
        (T["prepared_for"], _client_with_country),
        (T["prepared_by"],  seller_name or "—"),
        (T["date"],         today.strftime(_date_fmt)),
        (T["valid_until"],  valid_until.strftime(_date_fmt)),
    ]
    
    # Build metadata lines inside the cell
    for i, (label, value) in enumerate(rows_data):
        if i > 0:
            meta_cell.add_paragraph()  # spacing between rows
        p = meta_cell.paragraphs[i if i == 0 else -1]
        _para_space(p, before=2, after=2)
        
        # Label (bold, small, gray)
        lbl_run = p.add_run(label + ": ")
        lbl_run.font.bold = True
        lbl_run.font.size = Pt(10)
        lbl_run.font.color.rgb = IBM_GRAY
        
        # Value (regular, darker)
        val_run = p.add_run(value)
        val_run.font.size = Pt(10)
        val_run.font.color.rgb = IBM_DARK
    
    # ═══ 6. FOOTER NOTE (small, gray, centered) ═══
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_space(footer_p, before=60, after=0)
    footer_run = footer_p.add_run("IBM Business Partner Confidential")
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x8D, 0x8D, 0x8D)
    footer_run.font.italic = True


# ---------------------------------------------------------------------------
# Executive Summary narrative
# ---------------------------------------------------------------------------

def _add_executive_summary_text(doc, project, model_info, client_name, T, num_systems: int = 1):
    _add_section_heading(doc, T["sec_exec"])

    model_name   = model_info.get("name", "IBM Storage Solution")
    is_hybrid    = project.get("is_hybrid", False)
    drives_count = project.get("drives_count", 0)
    drive_type   = project.get("drive_type", "FlashCore Module 5")
    usable_tib   = project.get("usable_tib", 0.0)
    effective_tb = project.get("effective_tb", 0.0)
    cache_gb     = project.get("cache_gb", 256)
    support_info = project.get("support_info") or {}
    support_name = support_info.get("name", "IBM Storage Expert Care")
    support_years = support_info.get("years", 5)
    client_str   = f"for {client_name} " if client_name else ""

    # Multi-system note — shown above body when num_systems > 1
    if num_systems > 1:
        note_p = doc.add_paragraph()
        _para_space(note_p, before=0, after=8)
        note_run = note_p.add_run(
            T["multi_system_note"].format(n=num_systems, model=model_name)
        )
        note_run.font.size = Pt(10)
        note_run.font.bold = True
        note_run.font.color.rgb = IBM_BLUE

    support_hours = T["support_24x7"] if support_info.get("fix_time") else T["support_9x5"]

    # All-flash models never have HDD
    _all_flash_models = {"FS7600", "FS9600", "FS9500", "FS7300", "FSC200"}
    _model_short_body = model_info.get("short", "")
    _has_hdd_body = (is_hybrid
                     and project.get("hdd_drives_count", 0) > 0
                     and _model_short_body not in _all_flash_models)

    # For hybrid configs with actual HDD: override the body to describe both tiers
    if _has_hdd_body:
        _hdd_cnt     = project.get("hdd_drives_count", 0)
        _hdd_dt      = project.get("hdd_drive_type", "NL-SAS HDD")
        _hdd_use_tib = project.get("hdd_usable_tib", 0.0)
        _nvme_use_tib= project.get("nvme_usable_tib", 0.0) or project.get("usable_tib", 0.0)
        _hdd_enc     = project.get("hdd_enclosure", "")
        _enc_str     = f" ({_hdd_enc} expansion enclosure)" if _hdd_enc else ""
        _is_en       = T.get("lang", "en") != "pl"
        if _is_en:
            body = (
                f"The proposed solution {client_str}is based on the {model_name} — "
                f"an enterprise hybrid NVMe + HDD storage system. "
                f"The configuration combines {drives_count} × {drive_type} (NVMe tier, "
                f"{_nvme_use_tib:.2f} TiB usable) with {_hdd_cnt} × {_hdd_dt} (HDD nearline tier, "
                f"{_hdd_use_tib:.2f} TiB usable{_enc_str}), "
                f"delivering {usable_tib:.2f} TiB total usable capacity.\n\n"
                f"The NVMe tier provides low-latency, high-throughput access for active workloads, "
                f"while the HDD nearline tier offers cost-effective high-capacity storage for "
                f"secondary data, backup targets, and tiered workloads — all within a single IBM Spectrum "
                f"Virtualize managed pool with {cache_gb} GB of cache memory.\n\n"
                f"The solution is covered by {support_name} ({support_years}-year term), providing "
                f"{support_hours}."
            )
        else:
            body = (
                f"Proponowane rozwiązanie {client_str}opiera się na macierzy {model_name} — "
                f"korporacyjnym hybrydowym systemie pamięci masowej NVMe + HDD. "
                f"Konfiguracja łączy {drives_count} × {drive_type} (warstwa NVMe, "
                f"{_nvme_use_tib:.2f} TiB pojemności użytecznej) z {_hdd_cnt} × {_hdd_dt} "
                f"(warstwa HDD nearline, {_hdd_use_tib:.2f} TiB użytecznej{_enc_str}), "
                f"zapewniając łącznie {usable_tib:.2f} TiB pojemności użytecznej.\n\n"
                f"Warstwa NVMe zapewnia szybki dostęp z niskim opóźnieniem dla aktywnych obciążeń, "
                f"natomiast warstwa HDD nearline oferuje pojemnościową przestrzeń ekonomiczną dla "
                f"danych wtórnych, kopii zapasowych i obciążeń tiered — zarządzanych w jednej "
                f"puli IBM Spectrum Virtualize z {cache_gb} GB pamięci podręcznej.\n\n"
                f"Rozwiązanie objęte jest pakietem {support_name} (umowa {support_years}-letnia), "
                f"zapewniającym {support_hours}."
            )
    else:  # all-flash or non-hybrid
        body = T["body"].format(
            client_str=client_str,
            model_name=model_name,
            drives_count=drives_count,
            drive_type=drive_type,
            usable_tib=usable_tib,
            effective_tb=effective_tb,
            cache_gb=cache_gb,
            support_name=support_name,
            support_years=support_years,
            support_hours=support_hours,
        )

    p = doc.add_paragraph(body)
    p.style.font.size = Pt(10)
    _para_space(p, before=0, after=12)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = IBM_DARK

    # Key highlights — use Polish list when lang=pl (highlights_pl), fall back to English
    _is_pl = T.get("lang") == "pl"
    highlights = list(model_info.get("highlights_pl" if _is_pl else "highlights", []))
    if not highlights:
        highlights = list(model_info.get("highlights", []))
    # Add FlashSystem Grid if not already present
    _grid_hl = ("FlashSystem Grid — scale-out architecture connecting up to 32 arrays into a single "
                "managed pool, enabling linear capacity and performance scaling without disruption."
                if T.get("lang") != "pl" else
                "FlashSystem Grid — architektura scale-out łącząca do 32 macierzy w jedną "
                "zarządzaną pulę, umożliwiająca liniowe skalowanie pojemności i wydajności.")
    if not any("grid" in h.lower() for h in highlights):
        highlights.append(_grid_hl)

    if highlights:
        hl_p = doc.add_paragraph()
        run = hl_p.add_run(T["key_highlights"])
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = IBM_BLUE
        _para_space(hl_p, before=6, after=2)

        for hl in highlights:
            bp = doc.add_paragraph(style="List Bullet")
            run = bp.add_run(hl)
            run.font.size = Pt(10)
            run.font.color.rgb = IBM_DARK
            _para_space(bp, before=0, after=1)

    # FlashSystem platform advantages — adv1/adv2 differ for SAS-based models (no FCM)
    _is_pl2    = T.get("lang") == "pl"
    _model_sht = model_info.get("short", "")
    _is_sas    = _model_sht in _SAS_MODELS

    adv_p = doc.add_paragraph()
    run = adv_p.add_run(T["platform_advantages"])
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = IBM_BLUE
    _para_space(adv_p, before=8, after=2)

    # Build advantage list — replace adv1/adv2 for SAS models (no FlashCore Modules)
    if _is_sas:
        _sas_adv = [
            (
                "WORM snapshots & software ransomware detection — immutable snapshot policies "
                "protect against ransomware encryption and accidental deletion; anomaly detection "
                "through IBM Storage Insights."
                if not _is_pl2 else
                "Migawki WORM i programowe wykrywanie ransomware — niezmienne polityki migawek "
                "chronią przed szyfrowaniem przez ransomware i przypadkowym usunięciem; "
                "wykrywanie anomalii przez IBM Storage Insights."
            ),
            (
                "Distributed RAID 6 with fast rebuild — dual drive failure tolerance with "
                "configurable rebuild areas; drives can be added online without disruption."
                if not _is_pl2 else
                "Distributed RAID 6 z szybką odbudową — tolerancja na 2 jednoczesne awarie "
                "napędów; obszary odbudowy konfigurowalne; napędy dodawalne online bez przestoju."
            ),
            T["adv3"],
            T["adv4"],
        ]
        adv_items = _sas_adv
    else:
        adv_items = [T[k] for k in ("adv1", "adv2", "adv3", "adv4")]

    for adv_text in adv_items:
        ap = doc.add_paragraph(style="List Bullet")
        run = ap.add_run(adv_text)
        run.font.size = Pt(10)
        run.font.color.rgb = IBM_DARK
        _para_space(ap, before=0, after=2)


# ---------------------------------------------------------------------------
# Configuration table
# ---------------------------------------------------------------------------

def _add_config_table(doc, project, model_info, T):
    model_code      = project.get("model_code", "")
    product_version = project.get("product_version", "")
    drives_count    = project.get("drives_count", 0)
    drive_type      = project.get("drive_type", "FlashCore Module 5")
    drive_feature   = project.get("drive_feature", "")
    cache_gb        = project.get("cache_gb", 256)
    io_groups       = project.get("io_groups", 1)
    enclosures      = project.get("enclosures", 1)
    encryption      = project.get("encryption", False)
    fc_ports        = project.get("fc_ports", 0)
    is_hybrid       = project.get("is_hybrid", False)

    # Drive label: "N × CODE — Description"
    # Priority: description from CSV (drive_type) > FEATURE_LABELS > feature code alone
    if drive_feature:
        # Show description (FCM type / drive name) — feature code goes in the row header only
        _desc = drive_type or get_feature_label(drive_feature)
        _drive_val = f"{drives_count} × {_desc}"
    else:
        _drive_val = f"{drives_count} × {drive_type}"

    # Drive slot availability
    _total_slots = model_info.get("total_drive_slots", 0)
    _free_slots  = _total_slots - drives_count if _total_slots > drives_count else 0

    # Models that never support HDD (all-flash only)
    _all_flash_models = {"FS7600", "FS9600", "FS9500", "FS7300", "FSC200"}
    _model_short = model_info.get("short", "")
    _has_hdd = (is_hybrid
                and project.get("hdd_drives_count", 0) > 0
                and _model_short not in _all_flash_models)

    rows = [
        (T["cfg_model"],      f"{model_info.get('name', '')} ({model_code})"),
        (T["cfg_fw"],         product_version or "—"),
        (T["cfg_iogroups"],   str(io_groups)),
        (T["cfg_enclosures"], str(enclosures)),
        (f"{T['cfg_drives']} ({drive_feature})" if drive_feature else T["cfg_drives"],
         _drive_val),
    ]
    if _total_slots:
        rows.append((
            T["cfg_drive_slots"],
            T["cfg_drive_slots_val"].format(total=_total_slots, used=drives_count, free=_free_slots),
        ))
    # FC port label: 16 Gb for SAS-based models (ALBG), 32 Gb for FCM-based models (ALB9)
    _is_sas_cfg = _model_short in _SAS_MODELS
    if fc_ports:
        if _is_sas_cfg:
            _fc_label = (f"{fc_ports} × 16 Gb Fibre Channel ports (ALBG)"
                         if T.get("lang") != "pl" else
                         f"{fc_ports} × 16 Gb Fibre Channel (ALBG)")
        else:
            _fc_label = T["cfg_host_fc"].format(n=fc_ports)
    else:
        _fc_label = T["cfg_host_eth"]

    rows += [
        (T["cfg_cache"],      T["cfg_cache_val"].format(n=cache_gb)),
        (T["cfg_host"],       _fc_label),
        (T["cfg_encryption"], T["cfg_enc_yes"] if encryption else T["cfg_enc_no"]),
    ]

    # HDD tier row — only when config actually has HDD drives
    if _has_hdd:
        _hdd_cnt  = project.get("hdd_drives_count", 0)
        _hdd_dt   = project.get("hdd_drive_type", "NL-SAS HDD")
        _hdd_feat = project.get("hdd_drive_feature", "")
        _hdd_enc  = project.get("hdd_enclosure", "")
        _hdd_lbl  = f"HDD Tier ({_hdd_feat})" if _hdd_feat else "HDD Tier"
        _hdd_val  = f"{_hdd_cnt} × {_hdd_dt}"
        if _hdd_enc:
            _hdd_val += f"  [{_hdd_enc} expansion]"
        rows.insert(5, (_hdd_lbl, _hdd_val))

    t = _make_two_col_table(doc, rows)
    return t


# ---------------------------------------------------------------------------
# Capacity table
# ---------------------------------------------------------------------------

def _add_capacity_table(doc, project, T):
    raw_tb        = project.get("raw_tb", 0.0)
    raw_tib       = project.get("raw_tib", 0.0)
    usable_tb     = project.get("usable_tb", 0.0)
    usable_tib    = project.get("usable_tib", 0.0)
    effective_tb  = project.get("effective_tb", 0.0)
    effective_tib = project.get("effective_tib", 0.0)
    recommended   = project.get("recommended_max_tib", 0.0)
    raid_type     = project.get("raid_type", "DRAID6")
    rebuild_areas = project.get("rebuild_areas", 1)
    is_hybrid     = project.get("is_hybrid", False)
    _raid_display = re.sub(r"^D(RAID\d+)$", r"\1 (Distributed)", raid_type)
    _is_pl        = T.get("lang") == "pl"

    # All-flash model check
    _all_flash_models = {"FS7600", "FS9600", "FS9500", "FS7300", "FSC200"}
    _mi = get_model_info(project.get("model_code", ""))
    _has_hdd = (is_hybrid
                and project.get("hdd_drives_count", 0) > 0
                and _mi.get("short", "") not in _all_flash_models)

    # Labels differ for hybrid: "Full Raw Space (SSD+HDD)" etc.
    _raw_lbl = ("Łączna poj. raw (SSD+HDD)" if _is_pl else "Full Raw Space (SSD+HDD)") if _has_hdd else T["cap_raw"]
    _use_lbl = ("Łączna poj. użyteczna (SSD+HDD)" if _is_pl else "Full Usable Space (SSD+HDD)") if _has_hdd else T["cap_usable"]

    rows = [
        (_raw_lbl,           f"{raw_tb:.2f} TB  /  {raw_tib:.2f} TiB"),
        (_use_lbl,           f"{usable_tb:.2f} TB  /  {usable_tib:.2f} TiB"),
        (T["cap_effective"], f"{effective_tb:.2f} TB  /  {effective_tib:.2f} TiB"),
        (T["cap_rec"],       T["cap_rec_val"].format(n=recommended) if recommended else "—"),
        (T["cap_raid"],      T["cap_raid_val"].format(raid=_raid_display, areas=rebuild_areas)),
        (T["cap_drives"],    T["cap_drives_val"].format(n=project.get("drives_count", 0))),
    ]

    # Per-tier breakdown — only when config has real HDD drives
    if _has_hdd:
        _nvme_raw  = project.get("nvme_raw_tb", 0.0)
        _nvme_rtib = project.get("nvme_raw_tib", 0.0)
        _nvme_use  = project.get("nvme_usable_tb", 0.0)
        _nvme_utib = project.get("nvme_usable_tib", 0.0)
        _nvme_n    = project.get("nvme_drives_count", 0) or project.get("drives_count", 0)
        _nvme_dt   = project.get("drive_type", "NVMe SSD")
        _hdd_raw   = project.get("hdd_raw_tb", 0.0)
        _hdd_rtib  = project.get("hdd_raw_tib", 0.0)
        _hdd_use   = project.get("hdd_usable_tb", 0.0)
        _hdd_utib  = project.get("hdd_usable_tib", 0.0)
        _hdd_n     = project.get("hdd_drives_count", 0)
        _hdd_dt    = project.get("hdd_drive_type", "NL-SAS HDD")
        _hdd_enc   = project.get("hdd_enclosure", "")
        _nvme_lbl  = "  NVMe" if not _is_pl else "  NVMe"
        _hdd_lbl   = "  HDD Nearline" if not _is_pl else "  HDD Nearline"
        _hdd_enc_s = f"  [{_hdd_enc} expansion]" if _hdd_enc else ""
        rows += [
            (_nvme_lbl + " drives",  f"{_nvme_n} × {_nvme_dt}"),
            (_nvme_lbl + " Raw",     f"{_nvme_raw:.2f} TB  /  {_nvme_rtib:.2f} TiB"),
            (_nvme_lbl + " Usable",  f"{_nvme_use:.2f} TB  /  {_nvme_utib:.2f} TiB"),
            (_hdd_lbl  + " drives",  f"{_hdd_n} × {_hdd_dt}{_hdd_enc_s}"),
            (_hdd_lbl  + " Raw",     f"{_hdd_raw:.2f} TB  /  {_hdd_rtib:.2f} TiB"),
            (_hdd_lbl  + " Usable",  f"{_hdd_use:.2f} TB  /  {_hdd_utib:.2f} TiB"),
        ]

    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Performance table
# ---------------------------------------------------------------------------

def _add_performance_table(doc, project, iops, T):
    iops_total = iops or project.get("perf_iops_total", 0)
    iops_read = project.get("perf_iops_read", 0)
    iops_write = project.get("perf_iops_write", 0)
    read_pct = project.get("perf_read_pct", 0.0)
    throughput = project.get("perf_throughput_mib", 0.0)
    latency = project.get("perf_latency_ms", 0.0)
    cache_hit = project.get("perf_cache_hit_pct", 0.0)
    block_size = project.get("perf_transfer_size_kib", 16)

    if not iops_total:
        p = doc.add_paragraph(T["no_perf_data"])
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = IBM_GRAY
        return

    iops_sub1 = project.get("perf_iops_max_sub1ms", 0)
    lat_sub1  = project.get("perf_latency_at_max_sub1ms", 0.0)

    iops_headline = iops_sub1 if iops_sub1 else iops_total
    iops_headline_label = (
        f"{iops_headline:,} IOPS  (max @ sub-1 ms / {lat_sub1:.3f} ms latency)"
        if iops_sub1 else f"{iops_headline:,} IOPS"
    )

    # Split iops_sub1 by read/write ratio for the per-direction rows
    _read_pct  = read_pct if read_pct else 70.0
    _write_pct = 100.0 - _read_pct
    if iops_sub1:
        _iops_r = round(iops_sub1 * _read_pct  / 100)
        _iops_w = round(iops_sub1 * _write_pct / 100)
    else:
        _iops_r = iops_read
        _iops_w = iops_write

    rows = [
        (T["perf_profile"],  T["perf_profile_val"].format(read=_read_pct, write=_write_pct, bs=block_size)),
        (T["perf_total"],    iops_headline_label),
        (T["perf_read"],     f"{_iops_r:,} IOPS ({_read_pct:.0f}%)" if _iops_r else "—"),
        (T["perf_write"],    f"{_iops_w:,} IOPS ({_write_pct:.0f}%)" if _iops_w else "—"),
        (T["perf_tp"],       T["perf_tp_val"].format(mib=throughput, gib=throughput/1024) if throughput else "—"),
        (T["perf_cache"],    f"{cache_hit:.1f}%" if cache_hit else "—"),
        (T["perf_note"],     T["perf_note_val"]),
    ]

    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Connectivity & data protection
# ---------------------------------------------------------------------------

def _add_connectivity_table(doc, project, model_info, T):
    fc_ports  = project.get("fc_ports", 0)
    cable_qty = project.get("cable_qty", 0)
    encryption = project.get("encryption", False)
    raid_type  = project.get("raid_type", "DRAID6")
    _is_pl     = T.get("lang") == "pl"
    _is_sas    = model_info.get("short", "") in _SAS_MODELS

    # FC port description: 16 Gb ALBG for SAS models, 32 Gb ALB9 for FCM models
    if fc_ports:
        if _is_sas:
            _fc_val = (f"{fc_ports} × 16 Gb FC ports (ALBG adapter pair)"
                       if not _is_pl else
                       f"{fc_ports} × 16 Gb FC (para adapterów ALBG)")
        else:
            _fc_val = T["conn_fc_val"].format(n=fc_ports)
    else:
        _fc_val = T["conn_fc_no"]

    # Internal bus: SAS bus for SAS-based models, PCIe 4.0 NVMe for FCM models
    _bus_val = (("SAS 12 Gb/s (drives to controllers)" if not _is_pl
                 else "SAS 12 Gb/s (napędy do kontrolerów)")
                if _is_sas else T["conn_bus_val"])

    # Ransomware detection: software-based (IBM Storage Insights) for SAS models
    _ransom_val = (("Software-based ransomware detection via IBM Storage Insights anomaly "
                    "detection (no per-drive hardware acceleration)"
                    if not _is_pl else
                    "Programowe wykrywanie ransomware przez IBM Storage Insights "
                    "(bez sprzętowego przyspieszenia per-napęd)")
                   if _is_sas else T["conn_ransom_val"])

    rows = [
        (T["conn_fc"],       _fc_val),
        (T["conn_eth"],      T["conn_eth_val"]),
        (T["conn_cables"],   T["conn_cables_val"].format(n=cable_qty) if cable_qty else T["conn_cables_no"]),
        (T["conn_mgmt"],     T["conn_mgmt_val"]),
        (T["conn_bus"],      _bus_val),
        (T["conn_drv_prot"], f"{raid_type} — tolerates 2 simultaneous drive failures"),
        (T["conn_enc"],      T["conn_enc_yes"] if encryption else T["conn_enc_no"]),
        (T["conn_ransom"],   _ransom_val),
    ]

    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Physical environment
# ---------------------------------------------------------------------------

def _add_environment_table(doc, project, T):
    rack_units    = project.get("rack_units", 1)
    power_typ     = project.get("power_kw_typical", 0.0)
    power_kva_typ = project.get("power_kva_typical", 0.0)
    power_max     = project.get("power_kw_max", 0.0)
    power_kva_max = project.get("power_kva_max", 0.0)
    cooling       = project.get("cooling_btu", 0.0)
    is_hybrid     = project.get("is_hybrid", False)

    # For hybrid: build a rack breakdown string e.g. "6 RU (1U ctrl + 5U 4662-92G)"
    if is_hybrid:
        _hdd_enc     = project.get("hdd_enclosure", "")
        _hdd_enc_qty = project.get("hdd_enclosure_qty", 1)
        _model_code  = project.get("model_code", "")
        _MODEL_RU    = {"5127": 1, "5078": 2, "5015": 2}
        _ENC_RU      = {"4662-92G": 5, "4662-92F": 5}
        _ctrl_ru     = _MODEL_RU.get(_model_code[:4], 1)
        _enc_ru      = _ENC_RU.get(_hdd_enc, 0)
        _enc_label   = f" + {_hdd_enc_qty} × {_hdd_enc} ({_enc_ru}U)" if _enc_ru and _hdd_enc else ""
        _rack_str    = f"{rack_units} RU  ({_ctrl_ru}U ctrl{_enc_label})"
    else:
        _rack_str    = T["env_rack_val"].format(n=rack_units)

    rows = [
        (T["env_rack"],    _rack_str),
        (T["env_pwr_typ"], T["env_pwr_val"].format(kw=power_typ, kva=power_kva_typ) if power_typ else "—"),
        (T["env_pwr_max"], T["env_pwr_val"].format(kw=power_max, kva=power_kva_max) if power_max else "—"),
        (T["env_cool"],    T["env_cool_val"].format(btu=cooling) if cooling else "—"),
        (T["env_psu"],     T["env_psu_val"]),
        (T["env_ff"],      T["env_ff_val"]),
    ]

    _make_two_col_table(doc, rows)


# ---------------------------------------------------------------------------
# Software & subscriptions
# ---------------------------------------------------------------------------

def _add_software_table(doc, project, T):
    features = project.get("features", [])
    sw_rows = []

    sw_map = {
        "5608-PC2": "IBM Storage Intelligence Control",
        "5608-B24": "5-Year Registration SWMA for IBM Storage Intelligence Control",
        "5775-STG": "IBM Support Line for Storage (5 Year)",
        "M3TDX6": "Per Storage Device SWMA 5 Year",
        "M3R6QE": "Per Storage Device — 1 Year SW Maintenance",
    }

    seen = set()
    for feat in features:
        code = feat.get("code", "")
        label = sw_map.get(code)
        if label and code not in seen:
            sw_rows.append((code, label))
            seen.add(code)

    if not sw_rows:
        sw_rows = [("5608-PC2", "IBM Storage Intelligence Control (included)")]

    rows = [(code, desc) for code, desc in sw_rows]
    _make_two_col_table(doc, rows, header=(T["sw_code"], T["sw_desc"]))


# ---------------------------------------------------------------------------
# Support section
# ---------------------------------------------------------------------------

def _add_support_section(doc, project, T):
    support_info = project.get("support_info") or {}
    if not support_info:
        p = doc.add_paragraph(T["no_support"])
        return

    level       = support_info.get("level", "—")
    name        = support_info.get("name", "—")
    years       = support_info.get("years", "—")
    coverage    = support_info.get("coverage", "—")
    fix_time    = support_info.get("fix_time", False)
    fix_hours   = support_info.get("fix_time_hours") or "committed fix-time"
    description = support_info.get("description", "")

    if fix_time:
        _fix_str = T["sup_fixtime_yes"].format(hours=fix_hours)
    else:
        _fix_str = T["sup_fixtime_no"]

    rows = [
        (T["sup_pkg"],     name),
        (T["sup_level"],   level),
        (T["sup_term"],    T["sup_years_unit"].format(n=years)),
        (T["sup_coverage"],coverage),
        (T["sup_fixtime"], _fix_str),
        (T["sup_desc"],    description),
    ]

    _make_two_col_table(doc, rows)

    # Comparison note
    note_p = doc.add_paragraph()
    _para_space(note_p, before=4, after=4)
    run = note_p.add_run(T["sup_comparison"])
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = IBM_GRAY
    run2 = note_p.add_run(T["sup_comparison_body"])
    run2.font.size = Pt(9)
    run2.font.color.rgb = IBM_GRAY


# ---------------------------------------------------------------------------
# Pricing table
# ---------------------------------------------------------------------------

def _add_pricing_table(doc, pricing, project, T):
    currency     = project.get("currency", "EUR")
    discount_pct = pricing["discount_pct"]
    num_systems  = pricing.get("num_systems", 1)
    today        = date.today()
    valid_until  = today + timedelta(days=30)
    _date_fmt    = "%d.%m.%Y" if T.get("date") == "Data" else "%B %d, %Y"

    # Info line
    info_p = doc.add_paragraph()
    _para_space(info_p, before=0, after=6)
    run = info_p.add_run(
        T["price_info"].format(
            pf=project.get("price_file_date", ""),
            d=discount_pct,
            vu=valid_until.strftime(_date_fmt),
        )
    )
    run.font.size = Pt(9)
    run.font.color.rgb = IBM_GRAY
    run.font.italic = True

    # Pricing table — 4 columns: Category | Qty | List Price | End User Price
    table = doc.add_table(rows=0, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.add_row()
    _mc = project.get("model_code", "")
    for i, text in enumerate((
        T["price_cat"],
        T["price_qty"],
        T["price_list"].format(curr=currency),
        T["price_eu_col"].format(curr=currency),
    )):
        cell = hdr.cells[i]
        _set_cell_bg(cell, IBM_BLUE)
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = IBM_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    line_items = [
        (T["price_hw"].format(mc=_mc),  pricing["list_hw"],      pricing["eu_hw"]),
        (T["price_sup"],                 pricing["list_support"], pricing["eu_support"]),
        (T["price_sw"],                  pricing["list_sw"],      pricing["eu_sw"]),
        (T["price_ship"],                pricing["shipping"],     pricing["eu_ship"]),
    ]

    for i, (label, list_val, eu_val) in enumerate(line_items):
        row = table.add_row()
        bg = RGBColor(0xF4, 0xF4, 0xF4) if i % 2 == 0 else IBM_WHITE
        for cell in row.cells:
            _set_cell_bg(cell, bg)

        # col 0 — category
        row.cells[0].paragraphs[0].add_run(label).font.size = Pt(9)
        # col 1 — quantity
        _set_cell_text(row.cells[1], str(num_systems), align=WD_ALIGN_PARAGRAPH.CENTER, size=9)
        # col 2 — list price
        _set_cell_text(row.cells[2], f"{list_val:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)
        # col 3 — end user price
        _set_cell_text(row.cells[3], f"{eu_val:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9)

    # Subtotal row (list prices)
    sub_row = table.add_row()
    for cell in sub_row.cells:
        _set_cell_bg(cell, IBM_DARK)
    run_lbl = sub_row.cells[0].paragraphs[0].add_run(T["price_total"])
    run_lbl.font.bold = True; run_lbl.font.size = Pt(9); run_lbl.font.color.rgb = IBM_WHITE
    _set_cell_text(sub_row.cells[1], str(num_systems), align=WD_ALIGN_PARAGRAPH.CENTER, size=9, bold=True, color=IBM_WHITE)
    _set_cell_text(sub_row.cells[2], f"{pricing['list_total']:,.2f}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, bold=True, color=IBM_WHITE)
    _set_cell_text(sub_row.cells[3], "—", align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, color=IBM_WHITE)

    # End User Price total row
    eu_row = table.add_row()
    EU_GREEN = RGBColor(0x19, 0x8A, 0x38)
    for cell in eu_row.cells:
        _set_cell_bg(cell, EU_GREEN)
    run_eu = eu_row.cells[0].paragraphs[0].add_run(T["price_eu_row"])
    run_eu.font.bold = True; run_eu.font.size = Pt(10); run_eu.font.color.rgb = IBM_WHITE
    _set_cell_text(eu_row.cells[1], str(num_systems), align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, color=IBM_WHITE)
    _set_cell_text(eu_row.cells[2], "—", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, bold=True, color=IBM_WHITE)
    _set_cell_text(eu_row.cells[3], f"{pricing['end_user_total']:,.2f} {currency}", align=WD_ALIGN_PARAGRAPH.RIGHT, size=11, bold=True, color=IBM_WHITE)

    # Footnote
    fn = doc.add_paragraph()
    _para_space(fn, before=4, after=0)
    run = fn.add_run(T["price_fn"])
    run.font.size = Pt(8)
    run.font.color.rgb = IBM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Next Steps
# ---------------------------------------------------------------------------

def _add_next_steps(doc, project, model_info: dict, client_name: str, seller_name: str, T):
    _client = client_name or "your organisation"
    steps = [
        (T["next_1_title"], T["next_1_body"].format(client=_client)),
        (T["next_2_title"], T["next_2_body"]),
        (T["next_3_title"], T["next_3_body"]),
        (T["next_4_title"], T["next_4_body"]),
    ]

    for i, (title, body) in enumerate(steps):
        row_p = doc.add_paragraph()
        _para_space(row_p, before=4 if i > 0 else 0, after=0)
        num_run = row_p.add_run(f"{i+1}. ")
        num_run.font.bold = True
        num_run.font.size = Pt(10)
        num_run.font.color.rgb = IBM_BLUE
        title_run = row_p.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(10)
        title_run.font.color.rgb = IBM_DARK

        body_p = doc.add_paragraph(body)
        body_p.style.font.size = Pt(9)
        _para_space(body_p, before=1, after=4)
        for run in body_p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = IBM_GRAY

    # Contact block
    contact_p = doc.add_paragraph()
    _para_space(contact_p, before=8, after=0)
    contact_run = contact_p.add_run(
        T["contact"].format(name=seller_name or "your IBM Sales Representative")
    )
    contact_run.font.size = Pt(9)
    contact_run.font.bold = True
    contact_run.font.color.rgb = IBM_BLUE

    # Documentation links block
    _short = model_info.get("short", "")
    _docs  = get_docs(_short) if _short else {}
    _docs_url = _docs.get("docs_url", "")
    _sm_url   = _docs.get("sales_manual_url", "")
    if _docs_url or _sm_url:
        doc.add_paragraph()
        heading_p = doc.add_paragraph()
        _para_space(heading_p, before=8, after=2)
        _model_label = model_info.get("name", "") or model_info.get("short", "")
        hr = heading_p.add_run(
            f"{T['docs_heading']} — {_model_label}" if _model_label else T["docs_heading"]
        )
        hr.font.size = Pt(9)
        hr.font.bold = True
        hr.font.color.rgb = IBM_DARK

        def _link_para(label: str, url: str) -> None:
            p = doc.add_paragraph()
            _para_space(p, before=1, after=1)
            # label in gray
            lbl = p.add_run(f"{label}: ")
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = IBM_GRAY
            # hyperlink run via OOXML
            r_id = p.part.relate_to(url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True)
            hl = OxmlElement("w:hyperlink")
            hl.set(qn("r:id"), r_id)
            wr = OxmlElement("w:r")
            rpr = OxmlElement("w:rPr")
            style = OxmlElement("w:rStyle")
            style.set(qn("w:val"), "Hyperlink")
            rpr.append(style)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "18")   # 9 pt = 18 half-points
            rpr.append(sz)
            wr.append(rpr)
            wt = OxmlElement("w:t")
            wt.text = url
            wr.append(wt)
            hl.append(wr)
            p._p.append(hl)

        if _docs_url:
            _link_para(T["docs_ibm_docs"], _docs_url)
        if _sm_url:
            _link_para(T["docs_sales_manual"], _sm_url)


# ---------------------------------------------------------------------------
# Footer disclaimer
# ---------------------------------------------------------------------------

def _add_footer_disclaimer(doc, T):
    doc.add_paragraph()
    _add_hrule(doc)
    p = doc.add_paragraph()
    _para_space(p, before=4, after=0)
    run = p.add_run(T["disclaimer"])
    run.font.size = Pt(8)
    run.font.color.rgb = IBM_GRAY
    run.font.italic = True


# ---------------------------------------------------------------------------
# Pricing calculation
# ---------------------------------------------------------------------------

def _calc_pricing(
    project: dict,
    discount_pct: float,
    num_systems: int = 1,
    eu_margin_pct: float = 15.0,
) -> dict:
    """
    Pricing model:
      EU price  = List × (1 - discount%)        ← discount gives EU price directly
      BP price  = EU × (1 - margin%)             ← BP is what the partner pays IBM
      List price = 100% reference (e-config)
    """
    d = discount_pct / 100.0
    n = max(1, int(num_systems))
    m = 1.0 - eu_margin_pct / 100.0

    list_hw      = project.get("list_price_hw", 0.0)
    list_sw      = project.get("list_price_sw", 0.0)
    list_support = project.get("list_price_support", 0.0)
    shipping     = project.get("shipping", 0.0)

    # EU price per line = list × (1 - discount%)
    eu_hw_1      = list_hw      * (1 - d)
    eu_sw_1      = list_sw      * (1 - d)
    eu_support_1 = list_support * (1 - d)
    eu_ship_1    = shipping                      # shipping non-discountable

    eu_total_1   = eu_hw_1 + eu_sw_1 + eu_support_1 + eu_ship_1

    # BP price per line = EU × (1 - margin%)
    bp_hw_1      = eu_hw_1      * m
    bp_sw_1      = eu_sw_1      * m
    bp_support_1 = eu_support_1 * m
    bp_ship_1    = eu_ship_1    * m

    bp_total_1   = bp_hw_1 + bp_sw_1 + bp_support_1 + bp_ship_1

    list_grand   = (list_hw + list_sw + list_support + shipping) * n

    return {
        "discount_pct":    discount_pct,
        "num_systems":     n,
        "eu_margin_pct":   eu_margin_pct,
        # list prices (scaled)
        "list_hw":         list_hw      * n,
        "list_sw":         list_sw      * n,
        "list_support":    list_support * n,
        "shipping":        shipping     * n,
        "list_total":      list_grand,
        # EU prices (scaled) — discount applied to list price
        "eu_hw":           eu_hw_1      * n,
        "eu_sw":           eu_sw_1      * n,
        "eu_support":      eu_support_1 * n,
        "eu_ship":         eu_ship_1    * n,
        "end_user_total":  eu_total_1   * n,
        # BP prices (scaled) — EU minus partner margin
        "bp_hw":           bp_hw_1      * n,
        "bp_sw":           bp_sw_1      * n,
        "bp_support":      bp_support_1 * n,
        "bp_total":        bp_total_1   * n,
    }


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _set_page_margins(doc):
    from docx.oxml.ns import qn
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def _set_default_font(doc):
    from docx.oxml.ns import qn
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)


def _add_logo_header(doc):
    logo_path = LOGOS_DIR / "ibm-logo-2026.png"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_space(p, before=0, after=4)
    run = p.add_run()
    if logo_path.exists():
        run.add_picture(str(logo_path), width=Cm(3))
    else:
        run.text = "IBM"
        run.font.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = IBM_BLUE


def _convert_to_png(src: Path, dst: Path) -> None:
    """Convert any Pillow-supported image to PNG."""
    try:
        from PIL import Image
        img = Image.open(str(src)).convert("RGBA")
        img.save(str(dst), "PNG")
    except Exception:
        pass


def _get_logo_png(svg_path: Path, png_path: Path) -> Path | None:
    """Convert SVG to PNG if needed. Returns path to usable image or None."""
    if png_path.exists():
        return png_path
    if not svg_path.exists():
        return None
    # Try cairosvg (best quality SVG renderer)
    try:
        import cairosvg
        cairosvg.svg2png(url=str(svg_path), write_to=str(png_path), output_width=400)
        return png_path
    except ImportError:
        pass
    # Try inkscape CLI
    try:
        import subprocess
        subprocess.run(
            ["inkscape", str(svg_path), "--export-filename", str(png_path), "--export-width=400"],
            capture_output=True, timeout=10, check=True
        )
        if png_path.exists():
            return png_path
    except Exception:
        pass
    # Last resort: render IBM text as PNG placeholder
    try:
        from PIL import Image, ImageDraw
        img = Image.new("RGBA", (400, 150), (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.rectangle([10, 10, 390, 140], fill=(0, 98, 255))
        draw.text((160, 50), "IBM", fill=(255, 255, 255))
        img.save(str(png_path), "PNG")
        return png_path
    except Exception:
        return None


def _add_section_heading(doc, text: str):
    p = doc.add_paragraph()
    _para_space(p, before=14, after=4)
    run = p.add_run(text.upper())
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = IBM_BLUE
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0062FF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _make_two_col_table(doc, rows, header=None):
    col_count = 2
    table = doc.add_table(rows=0, cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if header:
        hdr_row = table.add_row()
        for i, h in enumerate(header):
            _set_cell_bg(hdr_row.cells[i], IBM_BLUE)
            run = hdr_row.cells[i].paragraphs[0].add_run(h)
            run.font.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = IBM_WHITE

    for i, (label, value) in enumerate(rows):
        row = table.add_row()
        bg = RGBColor(0xF4, 0xF4, 0xF4) if i % 2 == 0 else IBM_WHITE

        lbl_cell = row.cells[0]
        val_cell = row.cells[1]
        lbl_cell.width = Cm(6)
        val_cell.width = Cm(10)
        _set_cell_bg(lbl_cell, bg)
        _set_cell_bg(val_cell, bg)

        lbl_run = lbl_cell.paragraphs[0].add_run(label)
        lbl_run.font.bold = True
        lbl_run.font.size = Pt(9)
        lbl_run.font.color.rgb = IBM_GRAY

        val_run = val_cell.paragraphs[0].add_run(value)
        val_run.font.size = Pt(9)
        val_run.font.color.rgb = IBM_DARK

    _para_space_after_table(doc)
    return table


def _para_space(p, before=0, after=6):
    from docx.oxml.ns import qn
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"), str(after * 20))
    pPr.append(spacing)


def _para_space_after_table(doc):
    p = doc.add_paragraph()
    _para_space(p, before=0, after=4)


def _set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), str(color))
    tcPr.append(shd)


def _set_cell_text(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, size=9, bold=False, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = IBM_DARK


def _add_hrule(doc):
    p = doc.add_paragraph()
    _para_space(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0062FF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    tblPr.append(tblBorders)
